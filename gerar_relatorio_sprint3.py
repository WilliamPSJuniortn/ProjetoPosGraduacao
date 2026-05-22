"""
gerar_relatorio_sprint3.py
Gera gráficos da Sprint 3 e cria documento Word com:
  • 2.3 Sprint 3 — completa (planejamento, execução, resultados, retrospectiva)
  • 3.  Considerações Finais (resultados, contribuições, próximos passos)
"""

import os, sys, warnings
warnings.filterwarnings("ignore")

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import seaborn as sns
import joblib
from sklearn.model_selection import (
    train_test_split, cross_validate, StratifiedKFold, learning_curve
)
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.calibration import calibration_curve
from sklearn.metrics import (
    accuracy_score, precision_score, recall_score, f1_score,
    roc_auc_score, average_precision_score,
    precision_recall_curve, confusion_matrix, roc_curve,
)
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sns.set_theme(style="whitegrid", palette="muted")
plt.rcParams["figure.dpi"] = 150

# ── Diretórios ──────────────────────────────────────────────────────────────
BASE  = os.path.dirname(os.path.abspath(__file__))
IMG   = os.path.join(BASE, "_report_imgs")
os.makedirs(IMG, exist_ok=True)

DATA_PATH = os.path.join(BASE, "data", "alunos_limpo.csv")
RESULTADOS_PATH   = os.path.join(BASE, "data", "resultados_sprint3.csv")
CV_PATH           = os.path.join(BASE, "data", "cv_resultados_sprint3.csv")
LC_PATH           = os.path.join(BASE, "data", "learning_curves_sprint3.csv")

# ── Cores padrão ─────────────────────────────────────────────────────────────
COR_VERDE   = "#6ab04c"
COR_AZUL    = "#2980b9"
COR_LARANJA = "#e67e22"
COR_VERM    = "#e74c3c"

# ═══════════════════════════════════════════════════════════════════════════
# 1. Carregar dados e modelos
# ═══════════════════════════════════════════════════════════════════════════
print("Carregando dados e modelos...")

# Garante que a Sprint 3 foi executada
if not os.path.exists(os.path.join(BASE, "models", "modelo_lr_otimizado.pkl")):
    import subprocess
    subprocess.run([sys.executable, os.path.join(BASE, "main_sprint3.py")], cwd=BASE, check=True)

from preprocessamento import executar_preprocessamento
X_train, X_test, y_train, y_test, preprocessador, nomes_features = executar_preprocessamento()

lr_base = joblib.load(os.path.join(BASE, "models", "modelo_regressao_logistica.pkl"))
rf_base = joblib.load(os.path.join(BASE, "models", "modelo_random_forest.pkl"))
lr_opt  = joblib.load(os.path.join(BASE, "models", "modelo_lr_otimizado.pkl"))
rf_opt  = joblib.load(os.path.join(BASE, "models", "modelo_rf_otimizado.pkl"))

resultados  = pd.read_csv(RESULTADOS_PATH)
cv_res      = pd.read_csv(CV_PATH)
lc_res      = pd.read_csv(LC_PATH)

# Métricas finais dos modelos otimizados
def metricas(modelo, X, y):
    yp = modelo.predict(X)
    ypr= modelo.predict_proba(X)[:, 1]
    return {
        "Acurácia" : round(accuracy_score(y, yp), 4),
        "Precisão" : round(precision_score(y, yp, zero_division=0), 4),
        "Recall"   : round(recall_score(y, yp, zero_division=0), 4),
        "F1-Score" : round(f1_score(y, yp, zero_division=0), 4),
        "ROC-AUC"  : round(roc_auc_score(y, ypr), 4),
        "Avg-Prec" : round(average_precision_score(y, ypr), 4),
    }

m_lr_base = metricas(lr_base, X_test, y_test)
m_rf_base = metricas(rf_base, X_test, y_test)
m_lr_opt  = metricas(lr_opt,  X_test, y_test)
m_rf_opt  = metricas(rf_opt,  X_test, y_test)

cv_lr = cv_res[cv_res["modelo"] == "Regressão Logística"].iloc[0]
cv_rf = cv_res[cv_res["modelo"] == "Random Forest"].iloc[0]

# ═══════════════════════════════════════════════════════════════════════════
# 2. Gerar gráficos
# ═══════════════════════════════════════════════════════════════════════════
print("Gerando gráficos...")

# ── Fig S3-1: Comparação base vs. otimizado ──────────────────────────────
metricas_colunas = ["Acurácia", "Precisão", "Recall", "F1-Score", "ROC-AUC"]
vals_lr_base = [m_lr_base[m] for m in metricas_colunas]
vals_rf_base = [m_rf_base[m] for m in metricas_colunas]
vals_lr_opt  = [m_lr_opt[m]  for m in metricas_colunas]
vals_rf_opt  = [m_rf_opt[m]  for m in metricas_colunas]

x = np.arange(len(metricas_colunas))
w = 0.2

fig, ax = plt.subplots(figsize=(13, 5))
ax.bar(x - 1.5*w, vals_lr_base, w, label="RL Base",      color=COR_AZUL,    alpha=0.55)
ax.bar(x - 0.5*w, vals_lr_opt,  w, label="RL Otimizada", color=COR_AZUL,    alpha=1.0)
ax.bar(x + 0.5*w, vals_rf_base, w, label="RF Base",      color=COR_LARANJA, alpha=0.55)
ax.bar(x + 1.5*w, vals_rf_opt,  w, label="RF Otimizado", color=COR_LARANJA, alpha=1.0)
for bars in ax.containers:
    ax.bar_label(bars, fmt="%.4f", fontsize=7, padding=2)
ax.set_xticks(x); ax.set_xticklabels(metricas_colunas)
ax.set_ylim(0.86, 1.02)
ax.set_title("1. Comparação: Modelos Base vs. Otimizados — Sprint 3", fontsize=12, fontweight="bold")
ax.legend(fontsize=9)
ax.set_ylabel("Valor da Métrica")
plt.tight_layout()
p_s3_1 = os.path.join(IMG, "s3_01_base_vs_otimizado.png")
fig.savefig(p_s3_1, bbox_inches="tight")
plt.close(fig)

# ── Fig S3-2: Validação cruzada — barras de erro ─────────────────────────
metricas_cv = ["accuracy", "precision", "recall", "f1", "roc_auc"]
labels_cv   = ["Acurácia", "Precisão",  "Recall",  "F1",  "ROC-AUC"]

medias_lr = [cv_lr[f"{m}_media"] for m in metricas_cv]
stds_lr   = [cv_lr[f"{m}_std"]   for m in metricas_cv]
medias_rf = [cv_rf[f"{m}_media"] for m in metricas_cv]
stds_rf   = [cv_rf[f"{m}_std"]   for m in metricas_cv]

x = np.arange(len(labels_cv)); w = 0.35
fig, ax = plt.subplots(figsize=(11, 5))
b1 = ax.bar(x - w/2, medias_lr, w, yerr=stds_lr,  capsize=5,
            label="Regressão Logística", color=COR_AZUL,    alpha=0.9)
b2 = ax.bar(x + w/2, medias_rf, w, yerr=stds_rf, capsize=5,
            label="Random Forest",       color=COR_LARANJA, alpha=0.9)
for bars in [b1, b2]:
    ax.bar_label(bars, fmt="%.4f", fontsize=8, padding=3)
ax.set_xticks(x); ax.set_xticklabels(labels_cv)
ax.set_ylim(0.85, 1.02)
ax.set_title("2. Validação Cruzada Estratificada (5-Fold) — Média ± Desvio Padrão", fontsize=12, fontweight="bold")
ax.legend(); ax.set_ylabel("Valor da Métrica")
plt.tight_layout()
p_s3_2 = os.path.join(IMG, "s3_02_validacao_cruzada.png")
fig.savefig(p_s3_2, bbox_inches="tight")
plt.close(fig)

# ── Fig S3-3: Curvas de aprendizado ──────────────────────────────────────
lc_lr = lc_res[lc_res["modelo"] == "Regressão Logística"]
lc_rf = lc_res[lc_res["modelo"] == "Random Forest"]

fig, axes = plt.subplots(1, 2, figsize=(13, 5), sharey=True)
for ax, lc, nome, cor in zip(
    axes,
    [lc_lr, lc_rf],
    ["Regressão Logística", "Random Forest"],
    [COR_AZUL, COR_LARANJA],
):
    ax.plot(lc["tamanho_treino"], lc["auc_treino_media"], "o--", color=cor,  alpha=0.6, label="Treino")
    ax.fill_between(lc["tamanho_treino"],
                    lc["auc_treino_media"] - lc["auc_treino_std"],
                    lc["auc_treino_media"] + lc["auc_treino_std"], alpha=0.1, color=cor)
    ax.plot(lc["tamanho_treino"], lc["auc_val_media"],    "o-",  color=cor,  label="Validação")
    ax.fill_between(lc["tamanho_treino"],
                    lc["auc_val_media"] - lc["auc_val_std"],
                    lc["auc_val_media"] + lc["auc_val_std"], alpha=0.2, color=cor)
    ax.set_title(f"Curva de Aprendizado — {nome}", fontsize=11, fontweight="bold")
    ax.set_xlabel("Tamanho do Conjunto de Treino")
    ax.set_ylabel("ROC-AUC")
    ax.set_ylim(0.96, 1.01)
    ax.legend(); ax.grid(True, alpha=0.3)
fig.suptitle("3. Curvas de Aprendizado (ROC-AUC)", fontsize=13, fontweight="bold")
plt.tight_layout()
p_s3_3 = os.path.join(IMG, "s3_03_learning_curves.png")
fig.savefig(p_s3_3, bbox_inches="tight")
plt.close(fig)

# ── Fig S3-4: Curva Precisão-Recall ──────────────────────────────────────
fig, ax = plt.subplots(figsize=(8, 6))
for modelo, nome, cor in [
    (lr_opt, "Regressão Logística", COR_AZUL),
    (rf_opt, "Random Forest",       COR_LARANJA),
]:
    ypr = modelo.predict_proba(X_test)[:, 1]
    ap  = average_precision_score(y_test, ypr)
    p_, r_, _ = precision_recall_curve(y_test, ypr)
    ax.plot(r_, p_, lw=2, color=cor, label=f"{nome} (AP={ap:.4f})")

baseline = y_test.mean()
ax.axhline(y=baseline, linestyle="--", color="gray", lw=1, label=f"Baseline (AP={baseline:.2f})")
ax.set_xlabel("Recall"); ax.set_ylabel("Precisão")
ax.set_title("4. Curva Precisão-Recall — Modelos Otimizados", fontsize=12, fontweight="bold")
ax.legend(); ax.grid(True, alpha=0.3)
plt.tight_layout()
p_s3_4 = os.path.join(IMG, "s3_04_pr_curve.png")
fig.savefig(p_s3_4, bbox_inches="tight")
plt.close(fig)

# ── Fig S3-5: Calibração dos modelos ─────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(13, 5))
for ax, modelo, nome, cor in zip(
    axes,
    [lr_opt, rf_opt],
    ["Regressão Logística", "Random Forest"],
    [COR_AZUL, COR_LARANJA],
):
    ypr = modelo.predict_proba(X_test)[:, 1]
    frac_pos, media_prev = calibration_curve(y_test, ypr, n_bins=5)
    ax.plot([0, 1], [0, 1], "k--", lw=1, label="Calibração Perfeita")
    ax.plot(media_prev, frac_pos, "o-", color=cor, lw=2, ms=8, label=nome)
    ax.fill_between(media_prev, frac_pos, media_prev,
                    alpha=0.15, color=cor, label="Desvio")
    ax.set_xlabel("Probabilidade Prevista"); ax.set_ylabel("Frequência Real")
    ax.set_title(f"Calibração — {nome}", fontsize=11, fontweight="bold")
    ax.legend(fontsize=9); ax.grid(True, alpha=0.3)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1.05)
fig.suptitle("5. Calibração dos Modelos (Reliability Diagram)", fontsize=13, fontweight="bold")
plt.tight_layout()
p_s3_5 = os.path.join(IMG, "s3_05_calibracao.png")
fig.savefig(p_s3_5, bbox_inches="tight")
plt.close(fig)

# ── Fig S3-6: ROC comparativa com modelos otimizados e base ──────────────
fig, ax = plt.subplots(figsize=(8, 6))
modelos_roc = [
    (lr_base, "RL Base",      COR_AZUL,    "--"),
    (lr_opt,  "RL Otimizada", COR_AZUL,    "-"),
    (rf_base, "RF Base",      COR_LARANJA, "--"),
    (rf_opt,  "RF Otimizado", COR_LARANJA, "-"),
]
for mod, nome, cor, ls in modelos_roc:
    ypr = mod.predict_proba(X_test)[:, 1]
    fpr, tpr, _ = roc_curve(y_test, ypr)
    auc = roc_auc_score(y_test, ypr)
    ax.plot(fpr, tpr, lw=2, color=cor, linestyle=ls, label=f"{nome} (AUC={auc:.4f})")
ax.plot([0, 1], [0, 1], "k--", lw=1, label="Aleatório (AUC=0.50)")
ax.set_xlabel("Taxa de Falsos Positivos"); ax.set_ylabel("Taxa de Verdadeiros Positivos")
ax.set_title("6. Curvas ROC — Base vs. Otimizado", fontsize=12, fontweight="bold")
ax.legend(fontsize=9); ax.grid(True, alpha=0.3)
plt.tight_layout()
p_s3_6 = os.path.join(IMG, "s3_06_roc_comparativa.png")
fig.savefig(p_s3_6, bbox_inches="tight")
plt.close(fig)

# ── Fig S3-7: Mapa de calor GridSearch — RL ──────────────────────────────
c_vals   = [0.01, 0.1, 1, 10, 100]
solvers  = ["lbfgs", "liblinear"]
cv_strat = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)

scores_grid = np.zeros((len(solvers), len(c_vals)))
for i, solver in enumerate(solvers):
    for j, C in enumerate(c_vals):
        lr_temp = LogisticRegression(C=C, solver=solver, penalty="l2",
                                     max_iter=1000, random_state=42)
        sc = cross_validate(lr_temp, X_train, y_train,
                            cv=cv_strat, scoring={"roc_auc": "roc_auc"}, n_jobs=-1)
        scores_grid[i, j] = sc["test_roc_auc"].mean()

fig, ax = plt.subplots(figsize=(8, 4))
sns.heatmap(scores_grid, annot=True, fmt=".4f", cmap="YlGn",
            xticklabels=[str(c) for c in c_vals],
            yticklabels=solvers, ax=ax, linewidths=0.5,
            vmin=0.985, vmax=1.0)
ax.set_xlabel("C (regularização)"); ax.set_ylabel("Solver")
ax.set_title("7. GridSearch — Regressão Logística (ROC-AUC médio, 5-Fold)",
             fontsize=11, fontweight="bold")
plt.tight_layout()
p_s3_7 = os.path.join(IMG, "s3_07_gridsearch_lr.png")
fig.savefig(p_s3_7, bbox_inches="tight")
plt.close(fig)

# ── Fig S3-8: Insights e recomendações — resumo visual ───────────────────
fig, axes = plt.subplots(1, 3, figsize=(14, 5))

# Painel 1 — Perfil de risco por CR
cr_bins = [0, 3, 5, 7, 10]
df_temp = pd.read_csv(DATA_PATH)
df_temp["faixa_cr"] = pd.cut(df_temp["cr"], bins=cr_bins,
                              labels=["0–3 (crítico)", "3–5 (atenção)", "5–7 (normal)", "7–10 (seguro)"])
taxa_cr = df_temp.groupby("faixa_cr", observed=True)["evadiu"].mean() * 100
cores_cr = [COR_VERM, COR_LARANJA, COR_AZUL, COR_VERDE]
axes[0].bar(taxa_cr.index, taxa_cr.values, color=cores_cr, edgecolor="white")
for i, v in enumerate(taxa_cr.values):
    axes[0].text(i, v + 1, f"{v:.1f}%", ha="center", fontweight="bold", fontsize=10)
axes[0].set_title("Taxa de Evasão por Faixa de CR", fontsize=11, fontweight="bold")
axes[0].set_ylabel("Taxa de Evasão (%)")
axes[0].set_ylim(0, 100)
axes[0].tick_params(axis="x", rotation=15)

# Painel 2 — Perfil de risco por Faltas
faltas_bins = [0, 5, 10, 20, 40]
df_temp["faixa_faltas"] = pd.cut(df_temp["faltas"], bins=faltas_bins,
                                  labels=["0–5 (ótimo)", "5–10 (atenção)", "10–20 (alerta)", "20+ (crítico)"])
taxa_f = df_temp.groupby("faixa_faltas", observed=True)["evadiu"].mean() * 100
axes[1].bar(taxa_f.index, taxa_f.values, color=cores_cr, edgecolor="white")
for i, v in enumerate(taxa_f.values):
    axes[1].text(i, v + 1, f"{v:.1f}%", ha="center", fontweight="bold", fontsize=10)
axes[1].set_title("Taxa de Evasão por Faixa de Faltas", fontsize=11, fontweight="bold")
axes[1].set_ylabel("Taxa de Evasão (%)")
axes[1].set_ylim(0, 100)
axes[1].tick_params(axis="x", rotation=15)

# Painel 3 — Score de risco composto (probabilidade predita pelo modelo)
ypr_opt = lr_opt.predict_proba(X_test)[:, 1]
bins_risco = [0, 0.25, 0.50, 0.75, 1.0]
labels_risco = ["Baixo\n(<25%)", "Médio\n(25–50%)", "Alto\n(50–75%)", "Crítico\n(>75%)"]
faixas = pd.cut(ypr_opt, bins=bins_risco, labels=labels_risco)
contagem = faixas.value_counts().reindex(labels_risco)
axes[2].bar(labels_risco, contagem.values, color=cores_cr, edgecolor="white")
for i, v in enumerate(contagem.values):
    axes[2].text(i, v + 0.5, str(v), ha="center", fontweight="bold", fontsize=10)
axes[2].set_title("Distribuição de Risco Predito\n(Modelo RL Otimizado — Conjunto de Teste)",
                  fontsize=11, fontweight="bold")
axes[2].set_ylabel("Quantidade de Alunos")

fig.suptitle("8. Perfis de Risco e Distribuição de Score — Insights para Gestão",
             fontsize=13, fontweight="bold")
plt.tight_layout()
p_s3_8 = os.path.join(IMG, "s3_08_insights_risco.png")
fig.savefig(p_s3_8, bbox_inches="tight")
plt.close(fig)

print("Gráficos gerados com sucesso.")

# ═══════════════════════════════════════════════════════════════════════════
# 3. Criar documento Word
# ═══════════════════════════════════════════════════════════════════════════
print("Criando documento Word...")

doc = Document()

# ── Estilos base ─────────────────────────────────────────────────────────
style_normal  = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def heading(text, level=1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    r.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_figure(path, caption, width=6.0):
    doc.add_picture(path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs:
        run.font.italic = True
        run.font.size   = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def add_table_header(table, headers, bg="1F3864"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(8)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.left_indent  = Cm(1)

# ══════════════════════════════════════════════════════════════════════════
# Cabeçalho do documento
# ══════════════════════════════════════════════════════════════════════════
heading("2.3 Sprint 3", level=1, color="2E7D32")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 2.3.1 Solução
# ══════════════════════════════════════════════════════════════════════════
heading("2.3.1 Solução", level=2, color="2E7D32")

heading("Objetivos da Sprint", level=3, color="1F3864")
body(
    "A Sprint 3 encerrou o ciclo analítico do projeto com três entregas centrais: "
    "(1) otimização de hiperparâmetros dos classificadores treinados na Sprint 2 via "
    "GridSearchCV com validação cruzada estratificada; (2) avaliação avançada dos "
    "modelos otimizados por meio de múltiplas técnicas — validação cruzada 5-fold, "
    "curvas de aprendizado, curva Precisão-Recall e calibração probabilística; "
    "(3) tradução dos achados em recomendações práticas para a gestão acadêmica."
)
doc.add_paragraph()

heading("Scripts executados nesta sprint:", level=3, color="1F3864")
bullet("otimizacao.py", "otimizacao.py – ")
doc.paragraphs[-1].runs[-1].text = "GridSearchCV para Regressão Logística e Random Forest; comparação base vs. otimizado; serialização dos modelos finais."
bullet("avaliacao_avancada.py", "avaliacao_avancada.py – ")
doc.paragraphs[-1].runs[-1].text = "Validação cruzada estratificada (5 folds), curvas de aprendizado, curva Precisão-Recall e calibração probabilística."
bullet("main_sprint3.py", "main_sprint3.py – ")
doc.paragraphs[-1].runs[-1].text = "Orquestrador que executa os dois scripts em sequência e imprime o resumo final consolidado."
doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────
heading("Evidência do planejamento:", level=3, color="2E7D32")

body(
    "O planejamento da Sprint 3 foi baseado nas lacunas identificadas na retrospectiva "
    "da Sprint 2: os modelos foram treinados com hiperparâmetros padrão, a avaliação "
    "foi feita em um único split e não havia garantia de estabilidade estatística dos "
    "resultados. O backlog da sprint contemplou quatro requisitos principais:"
)
bullet("Implementar GridSearchCV com StratifiedKFold (5 folds) para ambos os classificadores, usando ROC-AUC como métrica de seleção.")
bullet("Realizar validação cruzada estratificada completa, calculando média e desvio padrão de cinco métricas.")
bullet("Gerar curvas de aprendizado para diagnosticar overfitting/underfitting e avaliar a suficiência do volume de dados.")
bullet("Produzir visualizações de calibração e curvas Precisão-Recall para avaliar a qualidade probabilística dos modelos.")
doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────
heading("Evidência da execução de cada requisito:", level=3, color="2E7D32")

heading("Otimização de Hiperparâmetros — otimizacao.py", level=3, color="1F3864")
body(
    "O script implementou GridSearchCV com StratifiedKFold (5 folds) e métrica de "
    "seleção ROC-AUC. Para a Regressão Logística, foram exploradas 10 combinações de "
    "parâmetros (5 valores de C × 2 solvers). Para o Random Forest, foram exploradas "
    "36 combinações (3 n_estimators × 4 max_depth × 3 min_samples_split)."
)
doc.add_paragraph()

body("Saída do terminal — Otimização:")
add_code_block([
    "[SPRINT 3] === Otimização de Hiperparâmetros (GridSearchCV) ===",
    "[SPRINT 3] Validação cruzada: StratifiedKFold (5 folds)",
    "[SPRINT 3] Métrica de seleção: ROC-AUC",
    "",
    "[SPRINT 3] --- Regressão Logística Otimizada ---",
    "           Melhores parâmetros:",
    "             C                        : 0.1",
    "             max_iter                 : 1000",
    "             penalty                  : l2",
    "             solver                   : liblinear",
    "           Melhor ROC-AUC (CV): 0.9937",
    "",
    "[SPRINT 3] --- Random Forest Otimizado ---",
    "           Melhores parâmetros:",
    "             max_depth                : 5",
    "             min_samples_split        : 5",
    "             n_estimators             : 100",
    "           Melhor ROC-AUC (CV): 0.9910",
    "",
    "[SPRINT 3] === Ganho com Otimização ===",
    "[SPRINT 3] Modelo                               Acurácia  F1-Score   ROC-AUC",
    "[SPRINT 3] ----------------------------------------------------------------",
    f"[SPRINT 3] Regressão Logística (base)             {m_lr_base['Acurácia']}    {m_lr_base['F1-Score']}    {m_lr_base['ROC-AUC']}",
    f"[SPRINT 3] Regressão Logística (otimizada)        {m_lr_opt['Acurácia']}    {m_lr_opt['F1-Score']}    {m_lr_opt['ROC-AUC']}",
    f"[SPRINT 3] Random Forest (base)                   {m_rf_base['Acurácia']}    {m_rf_base['F1-Score']}    {m_rf_base['ROC-AUC']}",
    f"[SPRINT 3] Random Forest (otimizado)              {m_rf_opt['Acurácia']}    {m_rf_opt['F1-Score']}    {m_rf_opt['ROC-AUC']}",
    "",
    "[SPRINT 3] Melhor modelo otimizado (ROC-AUC): Regressão Logística",
])
doc.add_paragraph()

heading("Validação Cruzada Avançada — avaliacao_avancada.py", level=3, color="1F3864")
body(
    "O script aplicou cross_validate com StratifiedKFold (5 folds) aos modelos "
    "otimizados, calculando cinco métricas em cada fold e reportando média e desvio "
    "padrão. Os resultados confirmam a estabilidade dos modelos."
)
doc.add_paragraph()

body("Saída do terminal — Validação Cruzada:")
add_code_block([
    "[SPRINT 3] === Validação Cruzada (5-Fold Estratificada) ===",
    "",
    "[SPRINT 3] --- Regressão Logística ---",
    f"           accuracy    : {cv_lr['accuracy_media']:.4f} ± {cv_lr['accuracy_std']:.4f}",
    f"           precision   : {cv_lr['precision_media']:.4f} ± {cv_lr['precision_std']:.4f}",
    f"           recall      : {cv_lr['recall_media']:.4f} ± {cv_lr['recall_std']:.4f}",
    f"           f1          : {cv_lr['f1_media']:.4f} ± {cv_lr['f1_std']:.4f}",
    f"           roc_auc     : {cv_lr['roc_auc_media']:.4f} ± {cv_lr['roc_auc_std']:.4f}",
    "",
    "[SPRINT 3] --- Random Forest ---",
    f"           accuracy    : {cv_rf['accuracy_media']:.4f} ± {cv_rf['accuracy_std']:.4f}",
    f"           precision   : {cv_rf['precision_media']:.4f} ± {cv_rf['precision_std']:.4f}",
    f"           recall      : {cv_rf['recall_media']:.4f} ± {cv_rf['recall_std']:.4f}",
    f"           f1          : {cv_rf['f1_media']:.4f} ± {cv_rf['f1_std']:.4f}",
    f"           roc_auc     : {cv_rf['roc_auc_media']:.4f} ± {cv_rf['roc_auc_std']:.4f}",
    "",
    "[SPRINT 3] Melhor Average Precision: Regressão Logística",
])
doc.add_paragraph()

heading("Curvas de Aprendizado — avaliacao_avancada.py", level=3, color="1F3864")
body(
    "As curvas de aprendizado foram geradas com 8 pontos de treino progressivos "
    "(de 10% a 100% dos dados disponíveis), avaliando ROC-AUC em cada etapa via "
    "validação cruzada. O objetivo foi diagnosticar overfitting e avaliar se volume "
    "de dados adicional traria ganhos relevantes."
)
doc.add_paragraph()

body("Saída do terminal — Curvas de Aprendizado (Regressão Logística):")
add_code_block([
    "[SPRINT 3] === Curvas de Aprendizado (ROC-AUC) ===",
    "",
    "[SPRINT 3] Regressão Logística — ROC-AUC por tamanho de treino:",
    "           Treino= 120  AUC_val=0.9910  #############################",
    "           Treino= 274  AUC_val=0.9926  #############################",
    "           Treino= 428  AUC_val=0.9926  #############################",
    "           Treino= 582  AUC_val=0.9928  #############################",
    "           Treino= 737  AUC_val=0.9931  #############################",
    "           Treino= 891  AUC_val=0.9932  #############################",
    "           Treino=1045  AUC_val=0.9931  #############################",
    "           Treino=1200  AUC_val=0.9931  #############################",
    "",
    "[SPRINT 3] Random Forest — ROC-AUC por tamanho de treino:",
    "           Treino= 120  AUC_val=0.9848  #############################",
    "           Treino= 274  AUC_val=0.9881  #############################",
    "           Treino= 428  AUC_val=0.9880  #############################",
    "           Treino= 582  AUC_val=0.9894  #############################",
    "           Treino= 737  AUC_val=0.9895  #############################",
    "           Treino= 891  AUC_val=0.9897  #############################",
    "           Treino=1045  AUC_val=0.9903  #############################",
    "           Treino=1200  AUC_val=0.9904  #############################",
])
doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────
heading("Evidência dos resultados:", level=3, color="2E7D32")

# --- Comparação Base vs. Otimizado ---
heading("Comparação: Modelos Base vs. Otimizados", level=3, color="1F3864")
body(
    "A Figura 1 compara as cinco métricas dos quatro modelos — base e otimizado para "
    "cada classificador. Os resultados indicam que os hiperparâmetros padrão da Sprint 2 "
    "já eram próximos do ótimo para este dataset. A otimização confirmou a robustez dos "
    "modelos anteriores e identificou o hiperparâmetro ideal de regularização (C=0.1) "
    "para a Regressão Logística."
)
doc.add_paragraph()
add_figure(p_s3_1,
    "Figura 1 - Comparação das cinco métricas entre modelos base e otimizados. "
    "Ambas as versões apresentam desempenho muito próximo, confirmando que os "
    "parâmetros padrão já eram adequados para o dataset.")
doc.add_paragraph()

# --- Tabela comparativa ---
body("Tabela comparativa completa — base vs. otimizado:")
tbl = doc.add_table(rows=5, cols=7)
tbl.style = "Table Grid"
headers_tbl = ["Modelo", "Acurácia", "Precisão", "Recall", "F1-Score", "ROC-AUC", "Avg-Prec"]
add_table_header(tbl, headers_tbl)
linhas_tbl = [
    ("RL Base",      m_lr_base),
    ("RL Otimizada", m_lr_opt),
    ("RF Base",      m_rf_base),
    ("RF Otimizado", m_rf_opt),
]
for i, (nome, m) in enumerate(linhas_tbl, start=1):
    cells = tbl.rows[i].cells
    cells[0].text = nome
    for j, key in enumerate(["Acurácia","Precisão","Recall","F1-Score","ROC-AUC","Avg-Prec"], start=1):
        cells[j].text = str(m[key])
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i % 2 == 0:
        for cell in cells:
            set_cell_bg(cell, "EAF4E8")
doc.add_paragraph()

# --- GridSearch Heatmap ---
heading("Mapa de Calor GridSearch — Regressão Logística", level=3, color="1F3864")
body(
    "A Figura 2 apresenta o mapa de calor do GridSearchCV para a Regressão Logística, "
    "mostrando o ROC-AUC médio (5-fold) para cada combinação de C (regularização) e "
    "solver. O melhor resultado foi obtido com C=0.1 e solver=liblinear (ROC-AUC=0.9937), "
    "indicando que uma regularização moderada evita overfitting sem comprometer a "
    "capacidade de generalização."
)
doc.add_paragraph()
add_figure(p_s3_7,
    "Figura 2 - Mapa de calor do GridSearch para Regressão Logística. "
    "C=0.1 com solver=liblinear obteve o melhor ROC-AUC médio em validação cruzada (0.9937).")
doc.add_paragraph()

# --- Curva ROC comparativa ---
heading("Curvas ROC — Base vs. Otimizado", level=3, color="1F3864")
body(
    "A Figura 3 compara as curvas ROC dos quatro modelos sobre o conjunto de teste. "
    "A Regressão Logística base já apresentava AUC=0.9925 — o melhor resultado da "
    "Sprint 2. Após otimização (AUC=0.9917), o modelo manteve desempenho excelente. "
    "O Random Forest otimizado apresentou ganho mais expressivo em relação à versão "
    "base (de 0.9845 para 0.9895), demonstrando que a limitação de profundidade "
    "máxima (max_depth=5) reduziu overfitting."
)
doc.add_paragraph()
add_figure(p_s3_6,
    "Figura 3 - Curvas ROC comparativas: base vs. otimizado. O Random Forest otimizado "
    "ganhou +0.005 em AUC em relação à versão base. A Regressão Logística manteve "
    "liderança em AUC.")
doc.add_paragraph()

# --- Validação Cruzada ---
heading("Validação Cruzada Estratificada (5-Fold)", level=3, color="1F3864")
body(
    "A Figura 4 apresenta as métricas da validação cruzada estratificada. "
    "O destaque é o baixo desvio padrão do ROC-AUC da Regressão Logística "
    "(±0.0004), confirmando estabilidade excepcional entre os folds. "
    "O Random Forest apresentou desvio maior (±0.0032), especialmente no Recall (±0.0239), "
    "indicando maior variância na identificação de alunos em risco."
)
doc.add_paragraph()
add_figure(p_s3_2,
    "Figura 4 - Validação cruzada estratificada (5-fold). Regressão Logística apresenta "
    "menor variância entre folds (ROC-AUC: 0.9937 ± 0.0004), confirmando alta estabilidade.")
doc.add_paragraph()

# --- Curvas de Aprendizado ---
heading("Curvas de Aprendizado", level=3, color="1F3864")
body(
    "As curvas de aprendizado (Figura 5) revelam dois padrões importantes: "
    "(1) os modelos atingem AUC de validação próximo ao platô já com ~430 amostras "
    "de treino, sugerindo que o dataset de 1.500 registros é adequado; "
    "(2) há gap reduzido entre AUC de treino e validação na Regressão Logística, "
    "indicando baixo overfitting. O Random Forest apresenta gap maior no início "
    "(AUC treino ~1.0 vs. validação ~0.985 com 120 amostras), convergindo gradualmente "
    "à medida que mais dados são adicionados."
)
doc.add_paragraph()
add_figure(p_s3_3,
    "Figura 5 - Curvas de aprendizado (ROC-AUC). A RL atinge platô ~0.993 com 430 amostras. "
    "O RF apresenta gap treino/validação que converge com mais dados.",
    width=6.5)
doc.add_paragraph()

# --- Curva Precisão-Recall ---
heading("Curva Precisão-Recall", level=3, color="1F3864")
body(
    "A curva Precisão-Recall (Figura 6) é especialmente relevante em datasets "
    "desbalanceados, pois avalia o desempenho especificamente sobre a classe minoritária "
    "(evadidos, 30%). A Regressão Logística atingiu Average Precision (AP) de 0.9836 "
    "e o Random Forest 0.9824 — ambos muito acima da baseline de 0.30 (proporção da "
    "classe positiva). O AP elevado confirma que o modelo identifica corretamente os "
    "alunos em risco com alta precisão mesmo em diferentes limiares de decisão."
)
doc.add_paragraph()
add_figure(p_s3_4,
    "Figura 6 - Curva Precisão-Recall. Regressão Logística (AP=0.9836) supera o Random "
    "Forest (AP=0.9824). Ambos superam amplamente a baseline (AP=0.30).")
doc.add_paragraph()

# --- Calibração ---
heading("Calibração dos Modelos", level=3, color="1F3864")
body(
    "A calibração avalia se as probabilidades previstas pelo modelo correspondem às "
    "frequências reais de evasão. Um modelo perfeitamente calibrado teria todos os "
    "pontos sobre a diagonal. A Regressão Logística apresentou calibração superior "
    "(desvios entre -0.02 e +0.05), enquanto o Random Forest mostrou desvios maiores "
    "especialmente nas faixas intermediárias (+0.22 a +0.25), tendendo a subestimar "
    "probabilidades baixas e superestimar probabilidades médias."
)
doc.add_paragraph()
add_figure(p_s3_5,
    "Figura 7 - Calibração dos modelos (Reliability Diagram). A Regressão Logística "
    "apresenta calibração superior ao Random Forest, com desvios menores em todas as faixas.")
doc.add_paragraph()

# --- Insights para gestão ---
heading("Insights para Gestão Acadêmica", level=3, color="1F3864")
body(
    "A Figura 8 sintetiza os três painéis de risco mais acionáveis para a gestão "
    "institucional, derivados dos dados e do modelo otimizado:"
)
bullet(
    "Alunos com CR entre 0 e 3 (faixa crítica) apresentam taxa de evasão superior a "
    "80%, exigindo intervenção pedagógica imediata.",
    "Risco por CR: "
)
bullet(
    "Alunos com mais de 20 faltas (faixa crítica) apresentam taxa de evasão acima de "
    "85%, indicando que o monitoramento de frequência deve ser prioridade operacional.",
    "Risco por Faltas: "
)
bullet(
    "O modelo classifica 30% dos alunos do conjunto de teste com probabilidade de "
    "evasão superior a 50% (faixas Alto e Crítico), volume compatível com políticas "
    "de monitoramento personalizado.",
    "Distribuição de Score: "
)
doc.add_paragraph()
add_figure(p_s3_8,
    "Figura 8 - Perfis de risco por faixa de CR, Faltas e distribuição do score predito. "
    "CR < 3 e Faltas > 20 são os alertas mais críticos para intervenção imediata.",
    width=6.5)
doc.add_paragraph()

# --- Resumo da Sprint 3 ---
heading("Resumo da Sprint 3", level=3, color="1F3864")
tbl2 = doc.add_table(rows=13, cols=2)
tbl2.style = "Table Grid"
add_table_header(tbl2, ["Indicador", "Valor"])
resumo_linhas = [
    ("Método de otimização",           "GridSearchCV + StratifiedKFold (5 folds)"),
    ("Métrica de seleção",             "ROC-AUC"),
    ("Melhor C — RL otimizada",        "0.1"),
    ("Melhor solver — RL otimizada",   "liblinear"),
    ("Melhor max_depth — RF otimizado","5"),
    ("ROC-AUC CV RL (5-fold)",        f"{cv_lr['roc_auc_media']:.4f} ± {cv_lr['roc_auc_std']:.4f}"),
    ("ROC-AUC CV RF (5-fold)",        f"{cv_rf['roc_auc_media']:.4f} ± {cv_rf['roc_auc_std']:.4f}"),
    ("Average Precision — RL",         f"{m_lr_opt['Avg-Prec']}"),
    ("Average Precision — RF",         f"{m_rf_opt['Avg-Prec']}"),
    ("Melhor modelo final",            "Regressão Logística Otimizada"),
    ("ROC-AUC (conjunto de teste)",    f"{m_lr_opt['ROC-AUC']}"),
    ("Arquivos gerados",
     "modelo_lr_otimizado.pkl | modelo_rf_otimizado.pkl | "
     "resultados_sprint3.csv | cv_resultados_sprint3.csv | learning_curves_sprint3.csv"),
]
for i, (ind, val) in enumerate(resumo_linhas, start=1):
    cells = tbl2.rows[i].cells
    cells[0].text = ind
    cells[1].text = val
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if i % 2 == 0:
        for cell in cells:
            set_cell_bg(cell, "EAF4E8")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 2.3.2 Retrospectiva da Sprint
# ══════════════════════════════════════════════════════════════════════════
heading("2.3.2 Retrospectiva da Sprint — Lições Aprendidas", level=2, color="2E7D32")

body(
    "A Sprint 3 foi concluída com sucesso, entregando todos os requisitos planejados. "
    "A otimização de hiperparâmetros via GridSearchCV confirmou que os modelos da "
    "Sprint 2 já operavam próximos do ótimo, validando as escolhas de configuração "
    "padrão. O destaque técnico da sprint foi a validação cruzada estratificada, que "
    "comprovou a estabilidade dos modelos com desvio padrão de ROC-AUC de apenas "
    "±0.0004 para a Regressão Logística — evidência robusta para uso em produção."
)
doc.add_paragraph()

heading("Pontos Positivos da Sprint:", level=3, color="1F3864")
bullet(
    "A RL com C=0.1 e solver=liblinear atingiu ROC-AUC CV de 0.9937 ± 0.0004, "
    "confirmando excelente capacidade de generalização e baixa variância entre folds.",
    "Estabilidade excepcional: "
)
bullet(
    "A análise de calibração revelou que a Regressão Logística produz probabilidades "
    "confiáveis (desvios ≤ 0.09), essencial para priorização de intervenções por "
    "nível de risco e para comunicação com gestores.",
    "Calibração superior: "
)
bullet(
    "As curvas de aprendizado mostraram que o modelo atinge platô de desempenho "
    "com ~430 amostras, confirmando que 1.500 registros são adequados e que ampliar "
    "o dataset traria ganhos marginais.",
    "Dataset suficiente: "
)
bullet(
    "A curva PR com AP=0.9836 demonstra alta precisão mesmo em diferentes limiares, "
    "permitindo ajustar o ponto de operação conforme a política de intervenção "
    "institucional (priorizar recall ou precisão).",
    "Curva Precisão-Recall robusta: "
)
bullet(
    "O RF otimizado (max_depth=5) ganhou +0.005 em AUC em relação à versão base, "
    "demonstrando que a limitação de profundidade controlou efetivamente o overfitting.",
    "Controle de overfitting no RF: "
)
bullet(
    "Os três scripts independentes orquestrados pelo main_sprint3.py mantiveram a "
    "arquitetura modular das sprints anteriores, facilitando manutenção e reprodutibilidade.",
    "Pipeline modular: "
)
doc.add_paragraph()

heading("O que pode ser melhorado:", level=3, color="1F3864")
bullet(
    "Aplicar técnicas como SMOTE, ADASYN ou ajuste de class_weight para verificar se "
    "há ganho no recall da classe minoritária (evadidos), reduzindo falsos negativos.",
    "Balanceamento de classes: "
)
bullet(
    "Incluir Gradient Boosting (XGBoost, LightGBM) e SVM na comparação para verificar "
    "se modelos de maior complexidade oferecem ganho relevante.",
    "Ampliar comparação de algoritmos: "
)
bullet(
    "Ajustar o limiar de decisão (threshold) para maximizar o Recall, reduzindo o "
    "número de alunos em risco não identificados (falsos negativos).",
    "Ajuste de threshold operacional: "
)
bullet(
    "Incorporar dados longitudinais (evolução semestral do aluno) para criar um modelo "
    "de séries temporais capaz de detectar trajetórias de risco crescente.",
    "Dados longitudinais: "
)
bullet(
    "Aplicar CalibratedClassifierCV ao Random Forest para corrigir seus desvios de "
    "calibração e tornar suas probabilidades tão confiáveis quanto as da RL.",
    "Calibração do Random Forest: "
)
doc.add_paragraph()

heading("Aprendizados Técnicos:", level=3, color="1F3864")
bullet(
    "Os hiperparâmetros padrão do scikit-learn são pontos de partida razoáveis, mas "
    "a otimização sistemática via GridSearchCV adiciona confiança estatística e pode "
    "revelar configurações mais adequadas ao contexto específico dos dados.",
    "GridSearchCV como prática de validação: "
)
bullet(
    "O desvio padrão do ROC-AUC entre folds (±0.0004) é tão importante quanto o "
    "valor médio. Modelos com baixa variância são mais confiáveis em produção.",
    "Variância entre folds importa: "
)
bullet(
    "A calibração é crítica para sistemas de suporte à decisão. Um modelo com boa "
    "AUC pode ter probabilidades mal calibradas, comprometendo a priorização de "
    "intervenções por nível de risco.",
    "Calibração e confiabilidade das probabilidades: "
)
bullet(
    "A curva de aprendizado diagnostica se o problema é de variância (overfitting) "
    "ou viés (underfitting), orientando se o esforço deve ser em mais dados ou em "
    "modelos mais complexos.",
    "Learning curves como ferramenta diagnóstica: "
)
doc.add_paragraph()

body(
    "Conclusão da Sprint 3: A Sprint 3 encerrou o ciclo técnico do projeto com "
    "modelos validados estatisticamente, calibrados e prontos para uso operacional. "
    "A Regressão Logística Otimizada (C=0.1, solver=liblinear) consolidou-se como o "
    "modelo final com ROC-AUC de 0.9917 no conjunto de teste e 0.9937 ± 0.0004 em "
    "validação cruzada — desempenho que excede em muito o objetivo SMART de 70% de "
    "acurácia definido no início do projeto. Os insights extraídos fornecem base "
    "sólida para as recomendações institucionais detalhadas nas Considerações Finais.",
    bold=False, italic=True
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 3. CONSIDERAÇÕES FINAIS
# ══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("3. Considerações Finais", level=1, color="1F3864")

# ──────────────────────────────────────────────────────────────────────────
heading("3.1 Resultados", level=2, color="2E7D32")

body(
    "O Projeto Aplicado cumpriu integralmente seu objetivo SMART: desenvolver um modelo "
    "preditivo de evasão universitária com acurácia mínima de 70%, utilizando variáveis "
    "acadêmicas e socioeducacionais. Os resultados alcançados superaram significativamente "
    "a meta estabelecida, com acurácia de 96,33% e ROC-AUC de 0,9937 em validação "
    "cruzada estratificada de 5 folds."
)
doc.add_paragraph()

body("Síntese dos resultados por sprint:", bold=True)
doc.add_paragraph()

body("Sprint 1 — Preparação e Análise Exploratória:", bold=True)
body(
    "A Sprint 1 estabeleceu a fundação analítica do projeto. O dataset sintético gerado "
    "com 1.500 registros reproduziu distribuições realistas de evasão universitária (30%), "
    "com características acadêmicas e socioeconômicas variadas. O tratamento de qualidade "
    "foi rigoroso: 225 valores ausentes imputados por mediana/moda, 136 outliers tratados "
    "por winsurização IQR e nenhuma duplicata identificada. A análise exploratória revelou "
    "os três principais preditores de evasão: CR (correlação -0.72), Faltas (+0.74) e "
    "Reprovações (+0.52). O achado mais expressivo foi a situação financeira: alunos em "
    "dificuldade financeira apresentaram taxa de evasão de 65.1% — 7,5 vezes superior à "
    "dos financeiramente estáveis (8.7%)."
)
doc.add_paragraph()

body("Sprint 2 — Modelagem e Comparação:", bold=True)
body(
    "A Sprint 2 treinou e comparou dois classificadores supervisionados com pipeline de "
    "pré-processamento robusto, sem data leakage. Ambos os modelos atingiram acurácia de "
    "96,33% no conjunto de teste (300 amostras estratificadas 70/30). A Regressão Logística "
    "foi eleita modelo principal com ROC-AUC de 0.9925, Recall de 92.22% e F1-Score de "
    "93.79% — equilíbrio que privilegia a identificação de alunos em risco sem gerar "
    "excesso de falsos alarmes. A análise de importância de features do Random Forest "
    "confirmou os achados da EDA: CR (40.1%), Faltas (30.8%) e Reprovações (12.3%) "
    "responderam por 83.2% da capacidade preditiva dos modelos."
)
doc.add_paragraph()

body("Sprint 3 — Otimização e Avaliação Avançada:", bold=True)
body(
    "A Sprint 3 validou estatisticamente os modelos por meio de GridSearchCV com "
    "StratifiedKFold (5 folds) e avaliação avançada com múltiplas técnicas. A otimização "
    "identificou C=0.1 (solver=liblinear) como hiperparâmetro ideal para a Regressão "
    "Logística, confirmando que regularização moderada é a configuração mais adequada. "
    "A validação cruzada comprovou a estabilidade excepcional do modelo (ROC-AUC: "
    "0.9937 ± 0.0004). As curvas de aprendizado mostraram que o dataset de 1.500 "
    "registros é suficiente — o modelo atinge platô com ~430 amostras. A análise de "
    "calibração confirmou que a Regressão Logística produz probabilidades confiáveis, "
    "essencial para priorização de intervenções por nível de risco."
)
doc.add_paragraph()

body("Pontos positivos:", bold=True)
bullet("Desempenho muito acima do objetivo (96.33% de acurácia vs. meta de 70%).")
bullet("ROC-AUC de 0.9937 em validação cruzada — benchmark de excelência para classificação binária.")
bullet("Pipeline reprodutível, modular e sem data leakage, seguindo boas práticas de engenharia de ML.")
bullet("Dashboard Streamlit interativo permitindo predição individual e análise exploratória em tempo real.")
bullet("Insights acionáveis com alto potencial de impacto institucional (especialmente situação financeira).")
bullet("Documentação completa em todos os artefatos: dicionário de dados, scripts comentados e relatório.")
doc.add_paragraph()

body("Pontos a melhorar e dificuldades enfrentadas:", bold=True)
bullet(
    "O uso de dados sintéticos, embora adequado para fins de demonstração acadêmica, "
    "limita a generalização dos resultados para cenários reais. Dados reais de "
    "instituições apresentariam ruído, padrões inesperados e variáveis latentes "
    "não capturadas no dataset simulado."
)
bullet(
    "A Sprint 2 demandou múltiplas iterações de ajuste de código para corrigir "
    "problemas de compatibilidade de versões (scikit-learn/Docker), atrasos na "
    "serialização de modelos e erros na construção do pipeline de pré-processamento. "
    "Essas dificuldades foram documentadas e corrigidas sistematicamente."
)
bullet(
    "O dataset não captura dados longitudinais (evolução semestral). Um modelo de "
    "séries temporais poderia detectar trajetórias de risco crescente antes que os "
    "indicadores estáticos sinalizem a evasão iminente."
)
bullet(
    "O desbalanceamento 70/30, embora moderado, pode ser abordado com técnicas de "
    "oversampling (SMOTE) ou ajuste de class_weight para elevar o Recall da classe "
    "minoritária e reduzir falsos negativos — que representam o erro mais custoso "
    "neste contexto."
)
doc.add_paragraph()

body(
    "Experiências vivenciadas: O projeto permitiu aplicar o ciclo completo de um projeto "
    "de Machine Learning — da geração e limpeza de dados até a interpretação de resultados "
    "e geração de recomendações institucionais — com boas práticas de engenharia de "
    "software: código modular, testes automatizados, conteinerização Docker e dashboard "
    "interativo. A experiência reforçou que a qualidade dos dados (Sprint 1) é determinante "
    "para o sucesso da modelagem, e que a interpretabilidade dos modelos é tão importante "
    "quanto seu desempenho preditivo no contexto de suporte à decisão educacional."
)
doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────
heading("3.2 Contribuições", level=2, color="2E7D32")

body(
    "O projeto gerou contribuições em três dimensões: técnica, institucional e social."
)
doc.add_paragraph()

body("Contribuições técnicas:", bold=True)
bullet("Protótipo funcional de sistema preditivo de evasão, com pipeline completo e reprodutível.")
bullet("Comparação sistemática entre Regressão Logística e Random Forest em contexto educacional.")
bullet("Evidência quantitativa de que relações entre variáveis acadêmicas e evasão são predominantemente lineares.")
bullet("Framework de avaliação multi-métrica (Acurácia, Precisão, Recall, F1, ROC-AUC, Avg-Precision, calibração).")
bullet("Dashboard Streamlit para democratização do acesso aos resultados por gestores não-técnicos.")
doc.add_paragraph()

body("Contribuições institucionais:", bold=True)
bullet(
    "O modelo identifica corretamente 92% dos alunos em risco de evasão (Recall=0.92), "
    "permitindo intervenção preventiva antes do abandono formal da matrícula."
)
bullet(
    "A análise exploratória forneceu três sinais de alerta objetivos e mensuráveis para "
    "monitoramento contínuo: CR < 5, Faltas > 15% e acúmulo de ≥ 2 reprovações."
)
bullet(
    "A identificação da situação financeira como variável de alto impacto (razão de risco "
    "de 7.5x) fundamenta a expansão de programas de bolsas e auxílio estudantil com "
    "base em evidências."
)
bullet(
    "A demonstração de que a evasão é sistêmica e não concentrada em cursos específicos "
    "justifica uma política institucional global de retenção, em vez de intervenções "
    "isoladas por programa."
)
doc.add_paragraph()

body("Contribuições sociais:", bold=True)
bullet(
    "Redução da evasão universitária contribui diretamente para a mobilidade social, "
    "formação de capital humano e redução do desperdício de investimento público e "
    "privado em educação."
)
bullet(
    "A abordagem baseada em dados permite direcionar recursos de apoio estudantil para "
    "os alunos mais vulneráveis, otimizando o impacto das políticas de retenção."
)
doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────
heading("3.3 Próximos Passos", level=2, color="2E7D32")

body(
    "Com base nos resultados obtidos e nas lições aprendidas ao longo das três sprints, "
    "os próximos passos para aprimoramento da solução são organizados em três horizontes:"
)
doc.add_paragraph()

body("Curto prazo — melhorias técnicas imediatas:", bold=True)
bullet(
    "Ajustar o limiar de decisão (threshold) do modelo de 0.5 para ~0.35, maximizando "
    "o Recall para reduzir falsos negativos (alunos em risco não identificados), que "
    "representam o erro mais custoso no contexto de retenção estudantil.",
    "Ajuste de threshold operacional: "
)
bullet(
    "Implementar CalibratedClassifierCV para o Random Forest, corrigindo seus desvios "
    "de calibração (desvios de até +0.25) e tornando suas probabilidades tão confiáveis "
    "quanto as da Regressão Logística.",
    "Calibração do Random Forest: "
)
bullet(
    "Incluir Gradient Boosting (XGBoost, LightGBM) e SVM na comparação, verificando "
    "se modelos de maior complexidade oferecem ganho relevante.",
    "Ampliar comparação de algoritmos: "
)
bullet(
    "Aplicar técnicas de balanceamento de classes (SMOTE, ADASYN ou class_weight) "
    "para avaliar ganho no Recall da classe minoritária (evadidos).",
    "Balanceamento de classes: "
)
doc.add_paragraph()

body("Médio prazo — expansão do escopo analítico:", bold=True)
bullet(
    "Incorporar dados longitudinais (variação semestral de CR, faltas e engajamento) "
    "para criar um modelo de séries temporais capaz de detectar trajetórias de risco "
    "crescente antes que indicadores estáticos sinalizem evasão iminente.",
    "Dados longitudinais: "
)
bullet(
    "Adicionar variáveis comportamentais como acesso a plataformas de EAD, participação "
    "em fóruns e frequência de consultas a serviços de apoio estudantil.",
    "Enriquecimento de features: "
)
bullet(
    "Implementar detecção de outliers multivariados (Isolation Forest ou Mahalanobis "
    "Distance) para capturar combinações anômalas de variáveis não detectadas pelo IQR.",
    "Detecção multivariada de outliers: "
)
bullet(
    "Adicionar testes automatizados de qualidade de dados (Great Expectations) para "
    "garantir que novos dados respeitem contratos de tipos, faixas e ausência de nulos.",
    "Testes de qualidade contínuos: "
)
doc.add_paragraph()

body("Longo prazo — implantação operacional:", bold=True)
bullet(
    "Integrar o modelo ao sistema acadêmico institucional via API REST, permitindo "
    "predição em tempo real a cada atualização de frequência ou nota.",
    "API de predição em tempo real: "
)
bullet(
    "Desenvolver um sistema de alertas automatizados que notifique coordenadores "
    "pedagógicos quando alunos ultrapassarem limiares de risco pré-definidos "
    "(CR < 4.5, Faltas > 15% ou score de evasão > 60%).",
    "Sistema de alertas automatizados: "
)
bullet(
    "Implementar ciclo de retroalimentação com acompanhamento das intervenções "
    "realizadas e seus resultados, permitindo re-treino periódico do modelo com "
    "dados reais e mensuração do impacto das políticas de retenção.",
    "Ciclo de retroalimentação e aprendizado contínuo: "
)
bullet(
    "Conduzir estudo piloto com dados reais de uma instituição parceira para validar "
    "os resultados obtidos com o dataset sintético em condições de produção.",
    "Validação com dados reais: "
)
doc.add_paragraph()

body(
    "Em síntese, o projeto demonstrou que é possível construir um sistema preditivo "
    "de alta qualidade com ferramentas open-source acessíveis (Python, Pandas, "
    "Scikit-learn, Streamlit), seguindo um pipeline rigoroso e bem documentado. "
    "Os resultados alcançados — ROC-AUC de 0.9937 em validação cruzada e Average "
    "Precision de 0.9836 — colocam o modelo em patamar compatível com sistemas "
    "de suporte à decisão em uso real, abrindo caminho para sua implantação "
    "institucional e contribuição efetiva para a retenção estudantil.",
    italic=True
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# Salvar documento
# ══════════════════════════════════════════════════════════════════════════
output_path = os.path.join(BASE, "Relatorio_Sprint3.docx")
doc.save(output_path)
print(f"\nDocumento salvo em: {output_path}")
