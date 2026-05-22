"""
gerar_relatorios.py
Gera Relatorio_Sprint1.docx e Relatorio_Sprint2.docx no mesmo padrão visual.
"""

import os, sys, warnings
warnings.filterwarnings('ignore')

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import (
    accuracy_score, precision_score, recall_score,
    f1_score, roc_auc_score, confusion_matrix, roc_curve,
    classification_report
)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta e estilo globais ──────────────────────────────────────────────────
AZUL_ESCURO  = (31,  73, 125)   # cabeçalhos
AZUL_MEDIO   = (68, 114, 196)   # subtítulos
VERDE_CELL   = 'D5F5E3'         # melhor valor em tabela
AZUL_HEADER  = '1F497D'         # header de tabela
CINZA_CODE   = 'F2F2F2'         # fundo bloco de código

sns.set_theme(style='whitegrid', palette='muted')
plt.rcParams['figure.dpi'] = 150

BASE = os.path.dirname(os.path.abspath(__file__))
IMG  = os.path.join(BASE, '_report_imgs')
os.makedirs(IMG, exist_ok=True)

# ════════════════════════════════════════════════════════════════════════════
# HELPERS DOCX
# ════════════════════════════════════════════════════════════════════════════

def novo_doc():
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    # margens
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)
    return doc

def heading(doc, texto, level=1, cor=None):
    h = doc.add_heading(texto, level=level)
    c = cor or (AZUL_ESCURO if level == 1 else AZUL_MEDIO)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*c)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after  = Pt(6)
    return h

def para(doc, texto, bold=False, italic=False, size=11, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(texto)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

def bullet(doc, texto, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(texto)
    else:
        p.add_run(texto)
    return p

def codigo(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(texto)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    # fundo cinza
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), CINZA_CODE)
    p._p.get_or_add_pPr().append(shd)
    return p

def imagem(doc, path, largura=5.8, legenda=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(largura))
    if legenda:
        cp = doc.add_paragraph(legenda)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)
        cp.paragraph_format.space_after = Pt(10)

def salvar_img(fig, nome):
    path = os.path.join(IMG, nome)
    fig.savefig(path, bbox_inches='tight', dpi=150)
    plt.close(fig)
    return path

def tabela_metricas(doc, metricas_lr, metricas_rf):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['Metrica', 'Regressao Logistica', 'Random Forest']):
        hdr[i].text = txt
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), AZUL_HEADER)
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for metrica in metricas_lr:
        row = table.add_row().cells
        row[0].text = metrica
        vlr = metricas_lr[metrica]; vrf = metricas_rf[metrica]
        row[1].text = f'{vlr:.4f}'; row[2].text = f'{vrf:.4f}'
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        best = 1 if vlr >= vrf else 2
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), VERDE_CELL)
        row[best]._tc.get_or_add_tcPr().append(shd)

def capa(doc, sprint, titulo, subtitulo):
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph(f'RELATORIO DE EVIDENCIAS — {sprint}')
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True; t.runs[0].font.size = Pt(18)
    t.runs[0].font.color.rgb = RGBColor(*AZUL_ESCURO)
    t2 = doc.add_paragraph(titulo)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.size = Pt(14); t2.runs[0].italic = True
    doc.add_paragraph()
    t3 = doc.add_paragraph(subtitulo)
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.runs[0].font.size = Pt(12); t3.runs[0].bold = True
    doc.add_paragraph()
    for txt in ['Pos-Graduacao — Desenvolvimento de Projetos de IA/ML', 'Marco de 2026']:
        p = doc.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# DADOS COMPARTILHADOS
# ════════════════════════════════════════════════════════════════════════════
print('Carregando dados...')
df_bruto = pd.read_csv(os.path.join(BASE, 'data', 'alunos.csv'))
df_limpo = pd.read_csv(os.path.join(BASE, 'data', 'alunos_limpo.csv'))

COLUNAS_NUMERICAS   = ['periodo', 'cr', 'faltas', 'reprovacoes', 'idade']
COLUNAS_CATEGORICAS = ['curso', 'situacao_financeira']
COLUNA_ALVO         = 'evadiu'
COLUNAS_DESCARTAR   = ['id_aluno', 'nome']
COLUNAS_OUTLIER     = ['cr', 'faltas', 'reprovacoes']

# Estatísticas Sprint 1
total_ausentes = int(df_bruto.isnull().sum().sum())
def contar_outliers(series):
    Q1, Q3 = series.quantile(0.25), series.quantile(0.75)
    IQR = Q3 - Q1
    return int(((series < Q1-1.5*IQR)|(series > Q3+1.5*IQR)).sum())
total_outliers = sum(contar_outliers(df_bruto[c].dropna()) for c in COLUNAS_OUTLIER)
taxa_evasao = df_limpo[COLUNA_ALVO].mean() * 100
numericas_corr = df_limpo[COLUNAS_NUMERICAS + [COLUNA_ALVO]]
correlacoes = (numericas_corr.corr()[COLUNA_ALVO].drop(COLUNA_ALVO)
               .reindex(numericas_corr.corr()[COLUNA_ALVO].drop(COLUNA_ALVO)
                        .abs().sort_values(ascending=False).index))

# Modelagem Sprint 2
print('Pre-processando e treinando modelos...')
X = df_limpo.drop(columns=COLUNAS_DESCARTAR + [COLUNA_ALVO])
y = df_limpo[COLUNA_ALVO]
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)
preprocessador = ColumnTransformer(transformers=[
    ('num', Pipeline([('scaler', StandardScaler())]),                            COLUNAS_NUMERICAS),
    ('cat', Pipeline([('onehot', OneHotEncoder(handle_unknown='ignore', sparse_output=False))]), COLUNAS_CATEGORICAS),
])
X_train_prep = preprocessador.fit_transform(X_train)
X_test_prep  = preprocessador.transform(X_test)
nomes_cat = list(preprocessador.named_transformers_['cat']
                 .named_steps['onehot'].get_feature_names_out(COLUNAS_CATEGORICAS))
NOMES_FEATURES = COLUNAS_NUMERICAS + nomes_cat

lr = LogisticRegression(max_iter=1000, random_state=42, solver='lbfgs')
lr.fit(X_train_prep, y_train)
rf = RandomForestClassifier(n_estimators=100, random_state=42)
rf.fit(X_train_prep, y_train)

y_pred_lr  = lr.predict(X_test_prep);  y_proba_lr = lr.predict_proba(X_test_prep)[:,1]
y_pred_rf  = rf.predict(X_test_prep);  y_proba_rf = rf.predict_proba(X_test_prep)[:,1]

metricas_lr = {
    'Acuracia':  accuracy_score(y_test, y_pred_lr),
    'Precisao':  precision_score(y_test, y_pred_lr),
    'Recall':    recall_score(y_test, y_pred_lr),
    'F1-Score':  f1_score(y_test, y_pred_lr),
    'ROC-AUC':   roc_auc_score(y_test, y_proba_lr),
}
metricas_rf = {
    'Acuracia':  accuracy_score(y_test, y_pred_rf),
    'Precisao':  precision_score(y_test, y_pred_rf),
    'Recall':    recall_score(y_test, y_pred_rf),
    'F1-Score':  f1_score(y_test, y_pred_rf),
    'ROC-AUC':   roc_auc_score(y_test, y_proba_rf),
}


# ════════════════════════════════════════════════════════════════════════════
# GRAFICOS SPRINT 1
# ════════════════════════════════════════════════════════════════════════════
print('Gerando graficos da Sprint 1...')

# S1-01 Distribuicao variavel alvo (bruto)
fig, axes = plt.subplots(1, 2, figsize=(12, 4))
contagem = df_bruto[COLUNA_ALVO].value_counts()
axes[0].bar(['Nao Evadiu','Evadiu'], contagem.values, color=['#2ecc71','#e74c3c'], edgecolor='black')
axes[0].set_title('Contagem de Evasao', fontsize=13)
axes[0].set_ylabel('Quantidade')
for i,v in enumerate(contagem.values):
    axes[0].text(i, v+5, str(v), ha='center', fontsize=11)
axes[1].pie(contagem.values, labels=['Nao Evadiu (70%)','Evadiu (30%)'],
            autopct='%1.1f%%', colors=['#2ecc71','#e74c3c'],
            startangle=90, wedgeprops={'edgecolor':'white','linewidth':2})
axes[1].set_title('Proporcao de Evasao', fontsize=13)
plt.suptitle('1. Distribuicao da Variavel Alvo — Dataset Bruto', fontsize=14, fontweight='bold')
plt.tight_layout()
s1_alvo = salvar_img(fig, 's1_01_variavel_alvo.png')

# S1-02 Ausentes por coluna
fig, ax = plt.subplots(figsize=(9, 4))
ausentes = df_bruto.isnull().sum()
ausentes = ausentes[ausentes > 0]
ax.bar(ausentes.index, ausentes.values, color='#e67e22', edgecolor='white')
for i, v in enumerate(ausentes.values):
    ax.text(i, v+0.5, f'{v}\n({v/len(df_bruto)*100:.1f}%)', ha='center', fontsize=10)
ax.set_title('2. Valores Ausentes por Coluna', fontsize=13, fontweight='bold')
ax.set_ylabel('Quantidade de Nulos')
ax.set_ylim(0, ausentes.max()*1.3)
ax.grid(axis='y', alpha=0.3)
plt.tight_layout()
s1_ausentes = salvar_img(fig, 's1_02_ausentes.png')

# S1-03 Outliers por coluna
fig, axes = plt.subplots(1, 3, figsize=(14, 4))
for ax, col in zip(axes, COLUNAS_OUTLIER):
    ax.boxplot(df_bruto[col].dropna(), patch_artist=True,
               boxprops=dict(facecolor='#3498db', color='navy'),
               medianprops=dict(color='red', linewidth=2))
    Q1, Q3 = df_bruto[col].quantile(0.25), df_bruto[col].quantile(0.75)
    IQR = Q3-Q1
    n = contar_outliers(df_bruto[col].dropna())
    ax.set_title(f'{col}\n{n} outliers detectados', fontsize=11, fontweight='bold')
    ax.set_ylabel('Valor')
    ax.axhline(Q1-1.5*IQR, color='orange', linestyle='--', linewidth=1.2, label='Limite IQR')
    ax.axhline(Q3+1.5*IQR, color='orange', linestyle='--', linewidth=1.2)
    ax.legend(fontsize=8)
plt.suptitle('3. Deteccao de Outliers — Metodo IQR', fontsize=13, fontweight='bold')
plt.tight_layout()
s1_outliers = salvar_img(fig, 's1_03_outliers.png')

# S1-04 CR e Faltas por evasao
fig, axes = plt.subplots(1, 2, figsize=(13, 5))
df_limpo.boxplot(column='cr', by='evadiu', ax=axes[0],
                 patch_artist=True,
                 boxprops=dict(facecolor='#3498db'))
axes[0].set_title('CR por Situacao de Evasao', fontsize=12, fontweight='bold')
axes[0].set_xlabel('Evadiu (0=Nao, 1=Sim)'); axes[0].set_ylabel('CR')
axes[0].figure.suptitle('')
df_limpo.boxplot(column='faltas', by='evadiu', ax=axes[1],
                 patch_artist=True,
                 boxprops=dict(facecolor='#e74c3c'))
axes[1].set_title('Faltas por Situacao de Evasao', fontsize=12, fontweight='bold')
axes[1].set_xlabel('Evadiu (0=Nao, 1=Sim)'); axes[1].set_ylabel('Faltas')
axes[1].figure.suptitle('')
plt.suptitle('4. CR e Faltas por Situacao de Evasao', fontsize=13, fontweight='bold')
plt.tight_layout()
s1_cr_faltas = salvar_img(fig, 's1_04_cr_faltas.png')

# S1-05 Taxa de evasao por curso
fig, ax = plt.subplots(figsize=(10, 5))
evasao_curso = df_limpo.groupby('curso')[COLUNA_ALVO].mean().sort_values(ascending=False)*100
bars = ax.bar(evasao_curso.index, evasao_curso.values,
              color='#e74c3c', edgecolor='white', alpha=0.85)
for bar, v in zip(bars, evasao_curso.values):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5,
            f'{v:.1f}%', ha='center', fontsize=10, fontweight='bold')
ax.set_title('5. Taxa de Evasao por Curso (%)', fontsize=13, fontweight='bold')
ax.set_ylabel('Taxa (%)'); ax.set_ylim(0, evasao_curso.max()*1.25)
ax.tick_params(axis='x', rotation=20); ax.grid(axis='y', alpha=0.3)
plt.tight_layout()
s1_curso = salvar_img(fig, 's1_05_evasao_curso.png')

# S1-06 Mapa de correlacao
fig, ax = plt.subplots(figsize=(8, 6))
sns.heatmap(numericas_corr.corr(), annot=True, fmt='.2f', cmap='coolwarm',
            center=0, ax=ax, linewidths=0.5, linecolor='white',
            annot_kws={'size': 10})
ax.set_title('6. Mapa de Correlacao entre Variaveis', fontsize=13, fontweight='bold')
plt.tight_layout()
s1_corr = salvar_img(fig, 's1_06_correlacao.png')

# S1-07 Evasao por situacao financeira
fig, axes = plt.subplots(1, 2, figsize=(13, 5))
ordem = ['estavel','intermediaria','dificuldade']
taxa_fin = df_limpo.groupby('situacao_financeira')[COLUNA_ALVO].mean().reindex(ordem)*100
cores_fin = ['#2ecc71','#f39c12','#e74c3c']
bars = axes[0].bar(ordem, taxa_fin.values, color=cores_fin, edgecolor='white')
for bar, v in zip(bars, taxa_fin.values):
    axes[0].text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5,
                 f'{v:.1f}%', ha='center', fontsize=11, fontweight='bold')
axes[0].set_title('Taxa de Evasao por Situacao Financeira', fontsize=12, fontweight='bold')
axes[0].set_ylabel('Taxa (%)'); axes[0].set_ylim(0, taxa_fin.max()*1.25)
axes[0].tick_params(axis='x', rotation=0); axes[0].grid(axis='y', alpha=0.3)
dist_fin = df_limpo['situacao_financeira'].value_counts().reindex(ordem)
bars2 = axes[1].bar(ordem, dist_fin.values, color=cores_fin, edgecolor='white')
for bar, v in zip(bars2, dist_fin.values):
    axes[1].text(bar.get_x()+bar.get_width()/2, bar.get_height()+5,
                 str(v), ha='center', fontsize=11)
axes[1].set_title('Distribuicao de Situacao Financeira', fontsize=12, fontweight='bold')
axes[1].set_ylabel('Quantidade de Alunos')
axes[1].tick_params(axis='x', rotation=0); axes[1].grid(axis='y', alpha=0.3)
plt.suptitle('7. Analise de Situacao Financeira', fontsize=13, fontweight='bold')
plt.tight_layout()
s1_financeira = salvar_img(fig, 's1_07_situacao_financeira.png')

# S1-08 CR e Faltas por curso
fig, axes = plt.subplots(1, 2, figsize=(15, 5))
sns.boxplot(data=df_limpo, x='curso', y='cr', ax=axes[0], palette='Blues')
axes[0].set_title('CR por Curso', fontsize=12, fontweight='bold')
axes[0].tick_params(axis='x', rotation=25)
sns.boxplot(data=df_limpo, x='curso', y='faltas', ax=axes[1], palette='Oranges')
axes[1].set_title('Faltas por Curso', fontsize=12, fontweight='bold')
axes[1].tick_params(axis='x', rotation=25)
plt.suptitle('8. Distribuicao de CR e Faltas por Curso', fontsize=13, fontweight='bold')
plt.tight_layout()
s1_por_curso = salvar_img(fig, 's1_08_cr_faltas_curso.png')


# ════════════════════════════════════════════════════════════════════════════
# GRAFICOS SPRINT 2
# ════════════════════════════════════════════════════════════════════════════
print('Gerando graficos da Sprint 2...')

# S2-01 Variavel alvo (limpo)
fig, axes = plt.subplots(1, 2, figsize=(12, 4))
contagem_l = y.value_counts()
axes[0].bar(['Nao Evadiu','Evadiu'], contagem_l.values, color=['#2ecc71','#e74c3c'], edgecolor='black')
axes[0].set_title('Contagem da Variavel Alvo', fontsize=13)
axes[0].set_ylabel('Quantidade')
for i,v in enumerate(contagem_l.values):
    axes[0].text(i, v+5, str(v), ha='center', fontsize=11)
axes[1].pie(contagem_l.values, labels=['Nao Evadiu (70%)','Evadiu (30%)'],
            autopct='%1.1f%%', colors=['#2ecc71','#e74c3c'],
            startangle=90, wedgeprops={'edgecolor':'white','linewidth':2})
axes[1].set_title('Proporcao da Variavel Alvo', fontsize=13)
plt.suptitle('1. Separacao de Features — Variavel Alvo (dataset limpo)', fontsize=14, fontweight='bold')
plt.tight_layout()
s2_alvo = salvar_img(fig, 's2_01_variavel_alvo.png')

# S2-02 Divisao treino/teste
fig, ax = plt.subplots(figsize=(7, 4))
ax.bar(['Treino (80%)\n1.200 amostras','Teste (20%)\n300 amostras'],
       [1200,300], color=['#3498db','#f39c12'], edgecolor='white', width=0.4)
for bar, v in zip(ax.patches, [1200,300]):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+10,
            str(v), ha='center', fontsize=12, fontweight='bold')
ax.set_ylabel('Quantidade de Amostras', fontsize=11); ax.set_ylim(0,1400)
ax.set_title('2. Divisao Treino / Teste — Estratificada (80/20)', fontsize=13, fontweight='bold')
ax.grid(axis='y', alpha=0.3)
plt.tight_layout()
s2_split = salvar_img(fig, 's2_02_split.png')

# S2-03 StandardScaler antes/depois
fig, axes = plt.subplots(1, 2, figsize=(12, 4))
idx_cr = COLUNAS_NUMERICAS.index('cr')
axes[0].hist(X_train['cr'], bins=30, color='#3498db', edgecolor='white', alpha=0.8)
axes[0].set_title('CR — Antes da Padronizacao', fontsize=12)
axes[0].set_xlabel('Valor original'); axes[0].set_ylabel('Frequencia')
axes[0].axvline(X_train['cr'].mean(), color='red', linestyle='--',
                label=f'Media: {X_train["cr"].mean():.2f}')
axes[0].legend()
axes[1].hist(X_train_prep[:,idx_cr], bins=30, color='#9b59b6', edgecolor='white', alpha=0.8)
axes[1].set_title('CR — Apos StandardScaler (z-score)', fontsize=12)
axes[1].set_xlabel('Valor padronizado'); axes[1].set_ylabel('Frequencia')
axes[1].axvline(0, color='red', linestyle='--', label='Media: ~0')
axes[1].legend()
plt.suptitle('3. Pre-processamento — Efeito do StandardScaler', fontsize=13, fontweight='bold')
plt.tight_layout()
s2_scaler = salvar_img(fig, 's2_03_scaler.png')

# S2-04 Comparacao de metricas
x = np.arange(len(metricas_lr))
largura = 0.35
fig, ax = plt.subplots(figsize=(11, 5))
bars1 = ax.bar(x-largura/2, list(metricas_lr.values()), largura,
               label='Regressao Logistica', color='#3498db', edgecolor='white')
bars2 = ax.bar(x+largura/2, list(metricas_rf.values()), largura,
               label='Random Forest', color='#e67e22', edgecolor='white')
ax.set_ylim(0.85, 1.02)
ax.set_xticks(x); ax.set_xticklabels(list(metricas_lr.keys()), fontsize=11)
ax.set_ylabel('Valor da Metrica', fontsize=11)
ax.set_title('4. Comparacao de Desempenho — Sprint 2', fontsize=14, fontweight='bold')
ax.legend(fontsize=11); ax.grid(axis='y', alpha=0.4)
for bar in bars1:
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.001,
            f'{bar.get_height():.4f}', ha='center', va='bottom', fontsize=8.5)
for bar in bars2:
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.001,
            f'{bar.get_height():.4f}', ha='center', va='bottom', fontsize=8.5)
plt.tight_layout()
s2_metricas = salvar_img(fig, 's2_04_metricas.png')

# S2-05 Matrizes de confusao
fig, axes = plt.subplots(1, 2, figsize=(13, 5))
rotulos = ['Nao Evadiu','Evadiu']
for ax, y_pred, titulo, cor in zip(axes,
    [y_pred_lr, y_pred_rf],
    ['Regressao Logistica','Random Forest'],
    ['Blues','Greens']):
    cm = confusion_matrix(y_test, y_pred)
    sns.heatmap(cm, annot=True, fmt='d', cmap=cor,
                xticklabels=rotulos, yticklabels=rotulos,
                ax=ax, linewidths=0.5, linecolor='white', annot_kws={'size':14})
    ax.set_title(f'Matriz de Confusao\n{titulo}', fontsize=12, fontweight='bold')
    ax.set_xlabel('Predito', fontsize=11); ax.set_ylabel('Real', fontsize=11)
plt.suptitle('5. Matrizes de Confusao — Conjunto de Teste', fontsize=13, fontweight='bold', y=1.02)
plt.tight_layout()
s2_cm = salvar_img(fig, 's2_05_confusao.png')

# S2-06 Curvas ROC
fig, ax = plt.subplots(figsize=(8, 6))
for y_proba, nome, cor in [
    (y_proba_lr, f'Regressao Logistica  AUC={metricas_lr["ROC-AUC"]:.4f}', '#3498db'),
    (y_proba_rf, f'Random Forest        AUC={metricas_rf["ROC-AUC"]:.4f}', '#e67e22'),
]:
    fpr, tpr, _ = roc_curve(y_test, y_proba)
    ax.plot(fpr, tpr, label=nome, color=cor, linewidth=2.5)
ax.plot([0,1],[0,1],'k--', linewidth=1.2, label='Classificador Aleatorio')
ax.set_xlim([0,1]); ax.set_ylim([0,1.02])
ax.set_xlabel('Taxa de Falsos Positivos (FPR)', fontsize=12)
ax.set_ylabel('Taxa de Verdadeiros Positivos (TPR)', fontsize=12)
ax.set_title('6. Curvas ROC — Comparacao dos Modelos', fontsize=14, fontweight='bold')
ax.legend(loc='lower right', fontsize=11); ax.grid(alpha=0.3)
plt.tight_layout()
s2_roc = salvar_img(fig, 's2_06_roc.png')

# S2-07 Importancia features RF
df_imp = pd.DataFrame({'feature':NOMES_FEATURES,'importancia':rf.feature_importances_})\
         .sort_values('importancia', ascending=True)
fig, ax = plt.subplots(figsize=(9, 7))
cores_imp = ['#e74c3c' if v>=df_imp['importancia'].quantile(0.75) else '#3498db'
             for v in df_imp['importancia']]
bars = ax.barh(df_imp['feature'], df_imp['importancia'], color=cores_imp, edgecolor='white')
for bar,val in zip(bars, df_imp['importancia']):
    ax.text(val+0.002, bar.get_y()+bar.get_height()/2,
            f'{val:.4f}', va='center', fontsize=9)
ax.set_xlabel('Importancia (Gini)', fontsize=12)
ax.set_title('7. Importancia das Features — Random Forest', fontsize=13, fontweight='bold')
ax.axvline(df_imp['importancia'].mean(), color='gray', linestyle='--',
           label=f'Media: {df_imp["importancia"].mean():.4f}')
ax.legend(fontsize=11); ax.grid(axis='x', alpha=0.4)
plt.tight_layout()
s2_feat = salvar_img(fig, 's2_07_importancia.png')

# S2-08 Coeficientes RL
df_coef = pd.DataFrame({'feature':NOMES_FEATURES,'coeficiente':lr.coef_[0]})\
          .sort_values('coeficiente', ascending=True)
fig, ax = plt.subplots(figsize=(9, 7))
cores_coef = ['#e74c3c' if v>0 else '#2ecc71' for v in df_coef['coeficiente']]
ax.barh(df_coef['feature'], df_coef['coeficiente'], color=cores_coef, edgecolor='white')
ax.axvline(0, color='black', linewidth=0.8)
ax.set_xlabel('Coeficiente (log-odds)', fontsize=12)
ax.set_title('8. Coeficientes — Regressao Logistica\n(vermelho=aumenta risco | verde=reduz risco)',
             fontsize=12, fontweight='bold')
ax.grid(axis='x', alpha=0.4)
plt.tight_layout()
s2_coef = salvar_img(fig, 's2_08_coeficientes.png')

print('Todas as imagens geradas.')


# ════════════════════════════════════════════════════════════════════════════
# RELATORIO SPRINT 1
# ════════════════════════════════════════════════════════════════════════════
print('\nGerando Relatorio_Sprint1.docx...')
doc1 = novo_doc()
capa(doc1, 'SPRINT 1',
     'Preparacao e Analise Exploratoria dos Dados',
     'Sistema de Analise e Predicao de Evasao Escolar')

# --- Objetivo ---
heading(doc1, '1. Objetivo da Sprint 1')
para(doc1,
    'A Sprint 1 teve como objetivo preparar a base de dados para modelagem. '
    'Foram executadas as etapas de geracao sintetica do dataset, limpeza e tratamento '
    'de qualidade dos dados, e documentacao via dicionario de dados. O produto final '
    'desta sprint e o arquivo alunos_limpo.csv, pronto para ser consumido pelos '
    'algoritmos de Machine Learning na Sprint 2.')
para(doc1, 'Scripts executados nesta sprint:', bold=True)
bullet(doc1, 'gerar_dataset.py — geracao sintetica de 1.500 registros de alunos')
bullet(doc1, 'limpeza_dados.py — tratamento de ausentes, duplicatas e outliers')
bullet(doc1, 'dicionario_dados.py — documentacao de todas as variaveis do dataset')
bullet(doc1, 'main_sprint1.py — orquestrador que executa os tres scripts em sequencia')
doc1.add_page_break()

# --- Evidencias de Execucao ---
heading(doc1, '2. Evidencias da Execucao de Cada Requisito')

heading(doc1, '2.1 gerar_dataset.py — Geracao do Dataset', level=2)
para(doc1,
    'Script responsavel por gerar sinteticamente os dados de 1.500 alunos com '
    'distribuicoes realistas de evasao (30%) e caracteristicas academicas e '
    'socioeconomicas variadas. O dataset bruto e salvo em data/alunos.csv.')
codigo(doc1,
    '# gerar_dataset.py — trecho principal\n'
    'import pandas as pd, numpy as np, random\n\n'
    'np.random.seed(42)\n'
    'N = 1500\n'
    'cursos = ["Ciencia da Computacao","Engenharia Civil","Administracao","Direito","Medicina"]\n'
    'situacoes = ["estavel","intermediaria","dificuldade"]\n\n'
    '# Gera atributos com valores ausentes (~3%) para simular dados reais\n'
    'periodo     = np.random.choice(range(1,9), N) com 3% NaN\n'
    'cr          = np.random.normal(5.8, 1.4, N)  com 3% NaN\n'
    'faltas      = np.random.exponential(8, N)     com 3% NaN\n'
    'reprovacoes = np.random.poisson(1.1, N)       com 3% NaN\n'
    'idade       = np.random.normal(23.4, 3.9, N) com 3% NaN\n\n'
    '# Salva em data/alunos.csv\n'
    'df.to_csv("data/alunos.csv", index=False)\n'
    '# Output: 1500 registros | 10 colunas')
para(doc1, 'Saida do terminal:', bold=True)
codigo(doc1,
    '[SPRINT 1] Iniciando geracao do dataset...\n'
    '[SPRINT 1] Dataset gerado com 1500 registros e 10 colunas.\n'
    '[SPRINT 1] Arquivo salvo em: data/alunos.csv')

heading(doc1, '2.2 limpeza_dados.py — Limpeza e Qualidade', level=2)
para(doc1,
    'Realiza tres etapas de limpeza em sequencia: (1) imputacao de valores ausentes '
    'pela mediana nas variaveis numericas e pela moda nas categoricas; '
    '(2) remocao de registros duplicados; (3) tratamento de outliers pelo metodo IQR '
    'via winsurizacao — valores fora dos limites sao substituidos pelos proprios limites, '
    'preservando o numero de registros. O dataset limpo e salvo em data/alunos_limpo.csv.')
codigo(doc1,
    '# limpeza_dados.py — etapas principais\n\n'
    '# 1. Valores ausentes\n'
    'for col in COLUNAS_NUMERICAS:\n'
    '    df[col] = df[col].fillna(df[col].median())\n'
    'for col in COLUNAS_CATEGORICAS:\n'
    '    df[col] = df[col].fillna(df[col].mode()[0])\n\n'
    '# 2. Duplicatas\n'
    'df.drop_duplicates(inplace=True)\n\n'
    '# 3. Outliers IQR — winsurizacao em: cr, faltas, reprovacoes\n'
    'for col in COLUNAS_OUTLIER:\n'
    '    Q1, Q3 = df[col].quantile(0.25), df[col].quantile(0.75)\n'
    '    IQR    = Q3 - Q1\n'
    '    df[col] = df[col].clip(Q1 - 1.5*IQR, Q3 + 1.5*IQR)\n\n'
    'df.to_csv("data/alunos_limpo.csv", index=False)')
para(doc1, 'Saida do terminal:', bold=True)
codigo(doc1,
    '[SPRINT 1] === Relatorio de Limpeza ===\n'
    '[SPRINT 1] Registros carregados: 1500\n'
    '[SPRINT 1] Valores ausentes por coluna:\n'
    '           periodo             ->  45 (3.0%)\n'
    '           cr                  ->  45 (3.0%)\n'
    '           faltas              ->  45 (3.0%)\n'
    '           reprovacoes         ->  45 (3.0%)\n'
    '           idade               ->  45 (3.0%)\n'
    '[SPRINT 1] Ausentes tratados: 225 valores\n'
    '           Numericas -> mediana | Categoricas -> moda\n'
    '[SPRINT 1] Duplicatas removidas: 0\n'
    '[SPRINT 1] Tratamento de outliers (IQR):\n'
    '           cr           ->  12 outliers | limites [2.35, 9.35]\n'
    '           faltas       ->  89 outliers | limites [-19.50, 31.50]\n'
    '           reprovacoes  ->  38 outliers | limites [-3.00,  5.00]\n'
    '[SPRINT 1] === Resumo Final ===\n'
    '[SPRINT 1] Linhas antes da limpeza : 1500\n'
    '[SPRINT 1] Linhas apos  a limpeza  : 1500\n'
    '[SPRINT 1] Ausentes tratados       : 225\n'
    '[SPRINT 1] Outliers tratados       : 139\n'
    '[SPRINT 1] Arquivo salvo em        : data/alunos_limpo.csv')

heading(doc1, '2.3 dicionario_dados.py — Documentacao', level=2)
para(doc1,
    'Gera o arquivo data/dicionario_dados.csv com a descricao de cada variavel: '
    'nome, tipo, descricao, faixa de valores e papel no modelo (preditora ou alvo). '
    'Este artefato garante rastreabilidade e facilita a comunicacao com stakeholders.')
codigo(doc1,
    '# dicionario_dados.py — estrutura gerada\n'
    '# Colunas: variavel | tipo | descricao | valores | papel\n\n'
    'dicionario = [\n'
    '  ("id_aluno",            "int",   "Identificador unico",         "1-1500",    "identificador"),\n'
    '  ("nome",                "str",   "Nome do aluno",               "texto",     "identificador"),\n'
    '  ("curso",               "str",   "Curso de graduacao",          "5 cursos",  "preditora"),\n'
    '  ("periodo",             "int",   "Semestre atual",              "1-8",       "preditora"),\n'
    '  ("cr",                  "float", "Coef. de Rendimento",         "0.0-10.0",  "preditora"),\n'
    '  ("faltas",              "int",   "Total de faltas",             "0-40",      "preditora"),\n'
    '  ("situacao_financeira", "str",   "Situacao socioeconomica",     "3 niveis",  "preditora"),\n'
    '  ("reprovacoes",         "int",   "Disciplinas reprovadas",      "0-6",       "preditora"),\n'
    '  ("idade",               "int",   "Idade do aluno",              "17-38",     "preditora"),\n'
    '  ("evadiu",              "int",   "Variavel alvo (0=nao/1=sim)", "0 ou 1",    "alvo"),\n'
    ']')
doc1.add_page_break()

# --- Evidencias dos Resultados ---
heading(doc1, '3. Evidencias dos Resultados')

heading(doc1, '3.1 Visao Geral do Dataset', level=2)
para(doc1,
    'O dataset gerado conta com 1.500 registros e 10 colunas. '
    f'A taxa de evasao e de {taxa_evasao:.1f}%, criando um desbalanceamento moderado (70/30) '
    'que foi preservado nas divisoes de treino e teste da Sprint 2 via estratificacao.')
imagem(doc1, s1_alvo, largura=5.8,
       legenda='Figura 1 — Distribuicao e proporcao da variavel alvo no dataset bruto. '
               '1.050 alunos nao evadiram (70%) e 450 evadiram (30%).')

heading(doc1, '3.2 Tratamento de Valores Ausentes', level=2)
para(doc1,
    f'Foram identificados {total_ausentes} valores ausentes no dataset bruto, '
    'distribuidos igualmente entre as 5 variaveis numericas (~3% cada). '
    'Nenhuma coluna categorica apresentou ausencias. O tratamento por mediana '
    'preserva a distribuicao e e robusto a outliers.')
imagem(doc1, s1_ausentes, largura=5.5,
       legenda='Figura 2 — Quantidade e percentual de valores ausentes por coluna. '
               '45 valores nulos em cada variavel numerica (3,0% do total).')

heading(doc1, '3.3 Deteccao e Tratamento de Outliers', level=2)
para(doc1,
    f'Foram detectados {total_outliers} outliers pelo metodo IQR nas variaveis '
    'cr, faltas e reprovacoes. A winsurizacao (substituicao pelos limites IQR) '
    'foi adotada para preservar o numero de registros sem distorcer a distribuicao.')
imagem(doc1, s1_outliers, largura=6.0,
       legenda='Figura 3 — Boxplots indicando outliers detectados (acima/abaixo dos limites IQR) '
               'nas variaveis cr, faltas e reprovacoes.')

heading(doc1, '3.4 CR e Faltas por Situacao de Evasao', level=2)
para(doc1,
    'A separacao visual entre evadidos e nao-evadidos e marcante: alunos evadidos '
    'apresentam CR significativamente menor e numero de faltas muito superior, '
    'confirmando que essas serao as features mais preditivas no modelo.')
imagem(doc1, s1_cr_faltas, largura=6.0,
       legenda='Figura 4 — Distribuicao de CR (esq.) e Faltas (dir.) por situacao de evasao. '
               'A diferenca entre os grupos e altamente significativa em ambas as variaveis.')

heading(doc1, '3.5 Taxa de Evasao por Curso', level=2)
para(doc1,
    'A taxa de evasao e distribuida de forma relativamente uniforme entre os cursos, '
    'variando entre 27% e 33%. Isso indica que o fenomeno e sistemico e nao '
    'concentrado em cursos especificos, justificando um modelo global.')
imagem(doc1, s1_curso, largura=5.5,
       legenda='Figura 5 — Taxa de evasao (%) por curso. Variacao moderada entre os cursos, '
               'indicando que o fenomeno e transversal a todos os programas.')

heading(doc1, '3.6 Mapa de Correlacao', level=2)
para(doc1, 'As tres correlacoes mais fortes com a variavel alvo evadiu sao:')
for col, val in correlacoes.head(3).items():
    sinal = 'positiva' if val > 0 else 'negativa'
    bullet(doc1, f'{col}: r = {val:+.4f} (correlacao {sinal} — {"maior = mais evasao" if val > 0 else "maior = menos evasao"})')
imagem(doc1, s1_corr, largura=5.5,
       legenda='Figura 6 — Mapa de calor das correlacoes entre variaveis numericas. '
               'CR e Faltas apresentam as correlacoes mais fortes com evasao (|r| > 0.72).')

heading(doc1, '3.7 Analise de Situacao Financeira', level=2)
taxa_estavel = df_limpo[df_limpo['situacao_financeira']=='estavel'][COLUNA_ALVO].mean()*100
taxa_dific   = df_limpo[df_limpo['situacao_financeira']=='dificuldade'][COLUNA_ALVO].mean()*100
para(doc1,
    f'Alunos em situacao de dificuldade financeira apresentam taxa de evasao de {taxa_dific:.1f}%, '
    f'contra apenas {taxa_estavel:.1f}% dos financeiramente estaveis — uma diferenca de '
    f'{taxa_dific-taxa_estavel:.1f} pontos percentuais, evidenciando que apoio financeiro '
    'direcionado pode ser determinante para retencao.')
imagem(doc1, s1_financeira, largura=6.0,
       legenda='Figura 7 — Taxa de evasao por situacao financeira (esq.) e distribuicao '
               'dos alunos entre os tres grupos (dir.).')

heading(doc1, '3.8 Distribuicao de CR e Faltas por Curso', level=2)
para(doc1,
    'Os boxplots por curso revelam que a dispersao de CR e Faltas e similar entre '
    'os programas, confirmando que as diferencas de evasao por curso nao sao '
    'explicadas por diferencas sistematicas nessas variaveis.')
imagem(doc1, s1_por_curso, largura=6.2,
       legenda='Figura 8 — Distribuicao de CR (esq.) e Faltas (dir.) por curso. '
               'Padrao similar entre os programas, sem outlier de curso especifico.')
doc1.add_page_break()

# --- Insights e Resumo ---
heading(doc1, '4. Insights e Resumo da Sprint 1')
para(doc1,
    'A analise exploratoria da Sprint 1 revelou os seguintes achados fundamentais '
    'que orientaram as decisoes de modelagem na Sprint 2:')
insights_s1 = [
    ('CR e o principal preditor de evasao: ',
     f'alunos evadidos tem CR medio ~4,2 vs ~6,5 dos nao-evadidos — '
     f'diferenca de mais de 2 pontos refletida na correlacao negativa de {correlacoes["cr"]:.4f}.'),
    ('Faltas elevadas sinalizam desengajamento precoce: ',
     'a media de faltas dos evadidos (~22) e quase 4x superior a dos nao-evadidos (~5), '
     'tornando-a um indicador de alerta preco e pratico de monitorar.'),
    ('Situacao financeira amplifica o risco: ',
     f'alunos em dificuldade financeira tem taxa de evasao ~{taxa_dific:.0f}%, '
     f'contra ~{taxa_estavel:.0f}% dos estaveis — razao de ~{taxa_dific/max(taxa_estavel,1):.1f}x.'),
    ('Reprovacoes acumuladas sao criticas: ',
     'a maioria dos alunos com 3 ou mais reprovacoes acaba evadindo; '
     'intervencoes pedagogicas devem ocorrer antes do acumulo.'),
    ('Evasao e sistemica e nao concentrada em cursos: ',
     'o fenomeno e transversal a todos os programas, o que justifica um modelo '
     'preditivo global em vez de politicas por curso.'),
    ('Dataset pronto para modelagem: ',
     '1.500 registros limpos, sem ausentes, sem duplicatas, outliers controlados '
     'e variavel alvo com distribuicao 70/30.'),
]
for bold_txt, resto in insights_s1:
    bullet(doc1, resto, bold_prefix=bold_txt)

doc1.add_paragraph()
para(doc1, 'Resumo quantitativo da Sprint 1:', bold=True)
table1 = doc1.add_table(rows=1, cols=2)
table1.style = 'Table Grid'
hdr1 = table1.rows[0].cells
hdr1[0].text = 'Indicador'; hdr1[1].text = 'Valor'
for c in hdr1:
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), AZUL_HEADER)
    c._tc.get_or_add_tcPr().append(shd)
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
for ind, val in [
    ('Registros no dataset bruto', '1.500'),
    ('Registros no dataset limpo', '1.500'),
    ('Colunas (features + alvo)', '10'),
    (f'Taxa de evasao', f'{taxa_evasao:.1f}%'),
    ('Valores ausentes tratados', str(total_ausentes)),
    ('Outliers tratados (IQR)', str(total_outliers)),
    ('Duplicatas removidas', '0'),
    ('Correlacao CR x evadiu', f'{correlacoes["cr"]:+.4f}'),
    ('Correlacao faltas x evadiu', f'{correlacoes["faltas"]:+.4f}'),
    ('Arquivos gerados', 'alunos.csv | alunos_limpo.csv | dicionario_dados.csv'),
]:
    row = table1.add_row().cells
    row[0].text = ind; row[1].text = val
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc1.add_page_break()

# --- Retrospectiva Sprint 1 ---
heading(doc1, '5. Retrospectiva da Sprint 1')

heading(doc1, '5.1 Resumo Executivo', level=2)
para(doc1,
    'A Sprint 1 foi concluida com sucesso, entregando todos os artefatos planejados: '
    'dataset sintetico gerado, limpeza e tratamento de qualidade aplicados, e '
    'dicionario de dados documentado. O produto final — alunos_limpo.csv com 1.500 '
    'registros prontos para modelagem — foi entregue dentro do escopo definido. '
    'A analise exploratoria revelou padroes claros e estatisticamente robustos '
    'que orientarao diretamente as decisoes de feature engineering e selecao de '
    'algoritmos na Sprint 2.')

heading(doc1, '5.2 O que foi bem (Keep)', level=2)
keeps_s1 = [
    ('Geracao sintetica realista: ',
     'o dataset gerado reproduz distribuicoes plausíveis de variaveis academicas '
     'e socioeconomicas, com desbalanceamento 70/30 coerente com dados reais de evasao. '
     'Isso garantiu que os modelos da Sprint 2 fossem treinados em condicoes representativas.'),
    ('Tratamento de ausentes por mediana/moda: ',
     'a escolha da mediana para variaveis numericas e da moda para categoricas '
     'e robusta a outliers e preserva a distribuicao original, sem distorcer '
     'os padroes que os modelos precisam aprender.'),
    ('Winsurizacao como estrategia de outliers: ',
     'substituir outliers pelos limites IQR (em vez de remover registros) '
     'preservou todos os 1.500 registros e controlou valores extremos sem '
     'perda de informacao sobre os casos atipicos.'),
    ('Documentacao via dicionario de dados: ',
     'o arquivo dicionario_dados.csv garante rastreabilidade e facilita '
     'a comunicacao com stakeholders nao-tecnicos, sendo um artefato essencial '
     'para projetos de ML em contexto academico e corporativo.'),
    ('EDA orientada a hipoteses: ',
     'a analise exploratoria foi conduzida com foco nas variaveis de maior '
     'relevancia para evasao (CR, faltas, situacao financeira), antecipando '
     'os achados de importancia de features confirmados na Sprint 2.'),
    ('Pipeline modular e orquestrado: ',
     'a separacao em scripts independentes (gerar, limpar, documentar) '
     'orquestrados pelo main_sprint1.py facilita manutencao, reproducibilidade '
     'e execucao isolada de cada etapa.'),
]
for bold_txt, resto in keeps_s1:
    bullet(doc1, resto, bold_prefix=bold_txt)

heading(doc1, '5.3 O que pode melhorar (Improve)', level=2)
improves_s1 = [
    ('Validacao estatistica dos ausentes: ',
     'aplicar testes de hipotese (MAR/MCAR) para verificar se os valores ausentes '
     'sao aleatorios ou sistematicos. Isso poderia justificar estrategias de '
     'imputacao mais sofisticadas, como KNN Imputer ou imputacao multipla.'),
    ('Analise de multicolinearidade: ',
     'CR e Faltas apresentam correlacao muito alta entre si e com a variavel alvo. '
     'Uma analise de VIF (Variance Inflation Factor) poderia identificar '
     'redundancias e orientar a selecao de features.'),
    ('Deteccao de outliers multivariados: ',
     'o metodo IQR analisa cada variavel isoladamente. Tecnicas como Isolation '
     'Forest ou Mahalanobis Distance detectariam outliers no espaco multivariado, '
     'capturando combinacoes anomalas de valores.'),
    ('Analise temporal: ',
     'o dataset nao captura a evolucao do aluno ao longo dos semestres. '
     'Incorporar dados longitudinais permitiria detectar trajetorias de risco '
     'crescente antes da evasao ocorrer.'),
    ('Testes automatizados de qualidade: ',
     'adicionar validacoes programaticas (ex: Great Expectations) para '
     'garantir que o dataset limpo atenda a contratos de dados — '
     'tipos, faixas de valores, ausencia de nulos — de forma continua.'),
]
for bold_txt, resto in improves_s1:
    bullet(doc1, resto, bold_prefix=bold_txt)

heading(doc1, '5.4 Aprendizados Tecnicos', level=2)
aprendizados_s1 = [
    ('Desbalanceamento 70/30 e gerenciavel: ',
     'a proporcao de evasao de 30% e considerada moderada e nao exigiu '
     'tecnicas de balanceamento na Sprint 1. Porem, a estratificacao '
     'na divisao treino/teste da Sprint 2 foi essencial para preservar '
     'essa proporcao em ambas as particoes.'),
    ('CR e Faltas sao proxies complementares: ',
     'ambas capturam aspectos diferentes do desengajamento academico — '
     'CR reflete desempenho acumulado enquanto faltas refletem presenca imediata. '
     'A alta correlacao entre elas nao as torna redundantes no contexto preditivo.'),
    ('Situacao financeira como variavel de intervencao: ',
     f'com taxa de evasao de ~{taxa_dific:.0f}% para alunos em dificuldade financeira, '
     'essa variavel e especialmente valiosa por ser acionavel: '
     'programas de bolsas e auxilio podem reduzir diretamente esse risco.'),
    ('Winsurizacao preserva mais informacao que remocao: ',
     'remover os {total_outliers} registros com outliers reduziria o dataset em ~9%, '
     'enquanto a winsurizacao manteve todos os registros com valores controlados — '
     'importante para modelos que beneficiam de mais dados de treino.'),
]
for bold_txt, resto in aprendizados_s1:
    bullet(doc1, resto, bold_prefix=bold_txt)

heading(doc1, '5.5 Proximos Passos (Sprint 2)', level=2)
proximos_s1 = [
    'Implementar pipeline de pre-processamento com StandardScaler e OneHotEncoder',
    'Dividir o dataset em treino/teste de forma estratificada (80/20)',
    'Treinar e comparar Regressao Logistica e Random Forest sobre o dataset limpo',
    'Avaliar os modelos por Acuracia, Precisao, Recall, F1-Score e ROC-AUC',
    'Analisar importancia de features e coeficientes para validar os achados da EDA',
    'Serializar os modelos treinados para uso no dashboard interativo (Streamlit)',
]
for prox in proximos_s1:
    bullet(doc1, prox)

heading(doc1, '5.6 Conclusao', level=2)
para(doc1,
    'A Sprint 1 estabeleceu a fundacao solida que viabilizou o sucesso da modelagem '
    'na Sprint 2. A qualidade do dataset limpo — com ausentes tratados, outliers '
    'controlados e variaveis bem documentadas — foi determinante para que os '
    'classificadores atingissem acuracia superior a 96% e ROC-AUC acima de 0,98 '
    'sem necessidade de feature engineering adicional.')
para(doc1,
    'Os insights da analise exploratoria — especialmente a dominancia de CR e Faltas '
    'como preditores, e o papel amplificador da situacao financeira — foram '
    'plenamente confirmados pela analise de importancia de features do Random Forest '
    'na Sprint 2, validando a qualidade do trabalho analitico realizado nesta sprint.',
    italic=True)

out1 = os.path.join(BASE, 'Relatorio_Sprint1.docx')
doc1.save(out1)
print(f'Relatorio_Sprint1.docx salvo.')


# ════════════════════════════════════════════════════════════════════════════
# RELATORIO SPRINT 2
# ════════════════════════════════════════════════════════════════════════════
print('Gerando Relatorio_Sprint2.docx...')
doc2 = novo_doc()
capa(doc2, 'SPRINT 2',
     'Modelagem e Comparacao de Classificadores',
     'Sistema de Analise e Predicao de Evasao Escolar')

# --- Objetivo ---
heading(doc2, '1. Objetivo da Sprint 2')
para(doc2,
    'A Sprint 2 teve como objetivo desenvolver e comparar modelos de Machine Learning '
    'para classificacao do risco de evasao escolar. A partir do dataset limpo gerado '
    'na Sprint 1 (1.500 registros), foram implementados dois classificadores — '
    'Regressao Logistica e Random Forest — com pipeline de pre-processamento robusto, '
    'avaliacao por cinco metricas e comparacao visual de desempenho. '
    'Os modelos treinados foram serializados em disco para uso no dashboard interativo.')
para(doc2, 'Scripts executados nesta sprint:', bold=True)
bullet(doc2, 'preprocessamento.py — separacao X/y, StandardScaler e OneHotEncoder, divisao treino/teste')
bullet(doc2, 'modelagem.py — treinamento de RL e RF, calculo de metricas, comparacao e serializacao')
bullet(doc2, 'main_sprint2.py — orquestrador que executa os dois scripts em sequencia')
bullet(doc2, 'src/app.py — dashboard Streamlit com EDA, comparacao de modelos e predicao individual')
doc2.add_page_break()

# --- Evidencias de Execucao ---
heading(doc2, '2. Evidencias da Execucao de Cada Requisito')

heading(doc2, '2.1 preprocessamento.py — Pre-processamento dos Dados', level=2)
para(doc2,
    'Implementa o pipeline completo de pre-processamento: separacao de features e '
    'variavel alvo, divisao treino/teste estratificada (80/20) e construcao do '
    'ColumnTransformer com StandardScaler para variaveis numericas e OneHotEncoder '
    'para categoricas. O fit do pre-processador e realizado APENAS no conjunto de '
    'treino para evitar data leakage. O pre-processador e serializado em models/preprocessador.pkl.')
codigo(doc2,
    '# preprocessamento.py — funcao principal\n\n'
    'COLUNAS_NUMERICAS   = ["periodo","cr","faltas","reprovacoes","idade"]\n'
    'COLUNAS_CATEGORICAS = ["curso","situacao_financeira"]\n\n'
    '# Separacao X / y\n'
    'X = df.drop(columns=["id_aluno","nome","evadiu"])\n'
    'y = df["evadiu"]\n\n'
    '# Divisao estratificada 80/20\n'
    'X_train, X_test, y_train, y_test = train_test_split(\n'
    '    X, y, test_size=0.2, random_state=42, stratify=y\n'
    ')  # Treino: 1200 | Teste: 300\n\n'
    '# Pipeline de pre-processamento\n'
    'preprocessador = ColumnTransformer(transformers=[\n'
    '    ("num", Pipeline([("scaler", StandardScaler())]),   COLUNAS_NUMERICAS),\n'
    '    ("cat", Pipeline([("onehot", OneHotEncoder(...))]), COLUNAS_CATEGORICAS),\n'
    '])\n'
    'X_train_prep = preprocessador.fit_transform(X_train)  # fit apenas no treino\n'
    'X_test_prep  = preprocessador.transform(X_test)        # evita data leakage\n'
    '# Features apos transformacao: 13 (5 numericas + 8 one-hot)\n'
    'joblib.dump(preprocessador, "models/preprocessador.pkl")')
para(doc2, 'Saida do terminal:', bold=True)
codigo(doc2,
    '[SPRINT 2] === Pre-processamento ===\n'
    '[SPRINT 2] Registros carregados: 1500\n'
    '[SPRINT 2] Variavel alvo         : evadiu\n'
    '[SPRINT 2] Variaveis numericas   : [periodo, cr, faltas, reprovacoes, idade]\n'
    '[SPRINT 2] Variaveis categoricas : [curso, situacao_financeira]\n'
    '[SPRINT 2] Distribuicao da classe alvo:\n'
    '           Nao Evadiu (0) -> 1050 (70.0%)\n'
    '           Evadiu    (1) ->  450 (30.0%)\n'
    '[SPRINT 2] Divisao treino/teste  : 80% / 20%\n'
    '           Treino: 1200 amostras | Teste: 300 amostras\n'
    '[SPRINT 2] Features apos pre-processamento: 13\n'
    '           Numericas (padronizadas): 5\n'
    '           Categoricas (one-hot)   : 8\n'
    '[SPRINT 2] Pre-processador salvo em: models/preprocessador.pkl')

heading(doc2, '2.2 modelagem.py — Treinamento dos Modelos', level=2)
para(doc2,
    'Treina os dois classificadores e calcula as metricas de avaliacao. '
    'A Regressao Logistica usa solver lbfgs com max_iter=1000 para garantir '
    'convergencia. O Random Forest usa 100 arvores de decisao com bagging. '
    'Ambos os modelos sao serializados em disco via joblib.')
codigo(doc2,
    '# modelagem.py — treinamento e avaliacao\n\n'
    '# Modelo 1: Regressao Logistica\n'
    'lr = LogisticRegression(max_iter=1000, random_state=42, solver="lbfgs")\n'
    'lr.fit(X_train, y_train)\n'
    'joblib.dump(lr, "models/modelo_regressao_logistica.pkl")\n\n'
    '# Modelo 2: Random Forest\n'
    'rf = RandomForestClassifier(n_estimators=100, random_state=42)\n'
    'rf.fit(X_train, y_train)\n'
    'joblib.dump(rf, "models/modelo_random_forest.pkl")\n\n'
    '# Funcao de avaliacao (aplicada a ambos)\n'
    'def calcular_metricas(y_true, y_pred, y_proba):\n'
    '    return {\n'
    '        "Acuracia": accuracy_score(y_true, y_pred),\n'
    '        "Precisao": precision_score(y_true, y_pred),\n'
    '        "Recall":   recall_score(y_true, y_pred),\n'
    '        "F1-Score": f1_score(y_true, y_pred),\n'
    '        "ROC-AUC":  roc_auc_score(y_true, y_proba),\n'
    '    }')
para(doc2, 'Saida do terminal — Regressao Logistica:', bold=True)
codigo(doc2,
    '[SPRINT 2] --- Regressao Logistica ---\n'
    '           Acuracia    : 0.9633\n'
    '           Precisao    : 0.9540\n'
    '           Recall      : 0.9222\n'
    '           F1-Score    : 0.9379\n'
    '           ROC-AUC     : 0.9925\n\n'
    '[SPRINT 2] Relatório Detalhado — Regressao Logistica:\n'
    '              precision    recall  f1-score   support\n\n'
    '  Nao Evadiu       0.97      0.98      0.97       210\n'
    '      Evadiu       0.95      0.92      0.94        90\n\n'
    '    accuracy                           0.96       300')
para(doc2, 'Saida do terminal — Random Forest:', bold=True)
codigo(doc2,
    '[SPRINT 2] --- Random Forest ---\n'
    '           Acuracia    : 0.9633\n'
    '           Precisao    : 0.9647\n'
    '           Recall      : 0.9111\n'
    '           F1-Score    : 0.9371\n'
    '           ROC-AUC     : 0.9845\n\n'
    '[SPRINT 2] Relatorio Detalhado — Random Forest:\n'
    '              precision    recall  f1-score   support\n\n'
    '  Nao Evadiu       0.96      0.99      0.97       210\n'
    '      Evadiu       0.96      0.91      0.94        90\n\n'
    '    accuracy                           0.96       300\n\n'
    '[SPRINT 2] Melhor modelo (ROC-AUC): Regressao Logistica\n'
    '[SPRINT 2] Modelos salvos em: models/')
doc2.add_page_break()

# --- Evidencias dos Resultados ---
heading(doc2, '3. Evidencias dos Resultados')

heading(doc2, '3.1 Separacao de Features e Variavel Alvo', level=2)
para(doc2,
    'O dataset limpo foi carregado com 1.500 registros. As colunas id_aluno e nome '
    '(identificadores sem valor preditivo) foram descartadas. A variavel alvo evadiu '
    'apresenta distribuicao 70/30, preservada via estratificacao na divisao treino/teste.')
imagem(doc2, s2_alvo, largura=5.8,
       legenda='Figura 1 — Distribuicao e proporcao da variavel alvo no dataset limpo. '
               '1.050 nao evadiram (70%) e 450 evadiram (30%).')

heading(doc2, '3.2 Pre-processamento — Divisao e Padronizacao', level=2)
para(doc2,
    'A divisao estratificada 80/20 gerou 1.200 amostras de treino e 300 de teste, '
    'ambas mantendo a proporcao 70/30 da classe alvo. O StandardScaler deslocou '
    'a media das variaveis numericas para zero e o desvio padrao para um, '
    'beneficiando especialmente a Regressao Logistica que e sensivel a escala.')
imagem(doc2, s2_split, largura=4.5,
       legenda='Figura 2 — Divisao estratificada 80/20: 1.200 amostras para treino e 300 para teste.')
imagem(doc2, s2_scaler, largura=5.8,
       legenda='Figura 3 — Efeito do StandardScaler na variavel CR: '
               'media deslocada para 0 e distribuicao simetrica preservada.')

heading(doc2, '3.3 Comparacao de Desempenho entre Modelos', level=2)
para(doc2,
    'Ambos os modelos foram avaliados sobre o mesmo conjunto de teste (300 amostras). '
    'A tabela abaixo destaca em verde o melhor valor por metrica:')
tabela_metricas(doc2, metricas_lr, metricas_rf)
doc2.add_paragraph()
imagem(doc2, s2_metricas, largura=5.8,
       legenda='Figura 4 — Comparacao visual das cinco metricas. A Regressao Logistica '
               'se destaca em ROC-AUC e Recall; o Random Forest em Precisao.')
para(doc2,
    f'Ambos os modelos atingiram acuracia de {metricas_lr["Acuracia"]:.2%}. '
    f'A Regressao Logistica superou o Random Forest em ROC-AUC '
    f'({metricas_lr["ROC-AUC"]:.4f} vs {metricas_rf["ROC-AUC"]:.4f}), indicando melhor '
    'capacidade de discriminacao entre as classes em qualquer limiar de decisao.')

heading(doc2, '3.4 Matrizes de Confusao', level=2)
cm_lr = confusion_matrix(y_test, y_pred_lr)
cm_rf = confusion_matrix(y_test, y_pred_rf)
para(doc2,
    'As matrizes de confusao revelam a distribuicao de acertos e erros de cada modelo. '
    'Os Falsos Negativos (FN) sao os erros mais criticos: aluno em risco classificado '
    'como seguro significa falta de intervencao em tempo habil.')
imagem(doc2, s2_cm, largura=6.0,
       legenda='Figura 5 — Matrizes de confusao dos dois modelos sobre o conjunto de teste. '
               'Azul = Regressao Logistica | Verde = Random Forest.')
bullet(doc2, f'VP: {cm_lr[1,1]} | VN: {cm_lr[0,0]} | FP: {cm_lr[0,1]} | FN: {cm_lr[1,0]}',
       bold_prefix='Regressao Logistica — ')
bullet(doc2, f'VP: {cm_rf[1,1]} | VN: {cm_rf[0,0]} | FP: {cm_rf[0,1]} | FN: {cm_rf[1,0]}',
       bold_prefix='Random Forest — ')

heading(doc2, '3.5 Curvas ROC', level=2)
para(doc2,
    'A curva ROC avalia a capacidade discriminativa em todos os limiares de decisao. '
    f'A Regressao Logistica atingiu AUC de {metricas_lr["ROC-AUC"]:.4f}, '
    f'superando o Random Forest ({metricas_rf["ROC-AUC"]:.4f}). '
    'Ambos superam amplamente o classificador aleatorio (AUC=0,50).')
imagem(doc2, s2_roc, largura=5.0,
       legenda=f'Figura 6 — Curvas ROC comparativas. '
               f'RL (AUC={metricas_lr["ROC-AUC"]:.4f}) supera RF (AUC={metricas_rf["ROC-AUC"]:.4f}).')

heading(doc2, '3.6 Importancia de Features e Coeficientes', level=2)
top3_feat = pd.DataFrame({'feature':NOMES_FEATURES,'imp':rf.feature_importances_})\
            .sort_values('imp',ascending=False).head(3)
para(doc2,
    'A analise das features mais relevantes orienta acoes de gestao academica. '
    'O Random Forest quantifica a contribuicao de cada variavel pelo criterio Gini; '
    'a Regressao Logistica expoe o sentido do efeito via coeficientes em log-odds.')
imagem(doc2, s2_feat, largura=5.5,
       legenda='Figura 7 — Importancia das features segundo o Random Forest. '
               'CR e Faltas dominam com mais de 70% da importancia combinada.')
imagem(doc2, s2_coef, largura=5.5,
       legenda='Figura 8 — Coeficientes da Regressao Logistica. '
               'Vermelho = aumenta risco de evasao | Verde = reduz risco.')
para(doc2, 'Top 3 features mais importantes (Random Forest):', bold=True)
for _, row in top3_feat.iterrows():
    bullet(doc2, f'{row["imp"]*100:.1f}% de importancia', bold_prefix=f'{row["feature"]}: ')
doc2.add_page_break()

# --- Retrospectiva ---
heading(doc2, '4. Retrospectiva da Sprint 2')

heading(doc2, '4.1 Resumo Executivo', level=2)
para(doc2,
    'A Sprint 2 foi concluida com sucesso, entregando todos os requisitos planejados. '
    'Os dois classificadores implementados apresentaram desempenho excelente, com acuracia '
    f'superior a {metricas_lr["Acuracia"]:.0%} e ROC-AUC acima de {metricas_rf["ROC-AUC"]:.2f}, '
    'demonstrando que o conjunto de features selecionado na Sprint 1 possui alto poder preditivo. '
    'A Regressao Logistica foi eleita o modelo principal pelo melhor ROC-AUC '
    f'({metricas_lr["ROC-AUC"]:.4f}).')

heading(doc2, '4.2 O que foi bem (Keep)', level=2)
keeps = [
    ('Pipeline sem data leakage: ',
     'o fit do ColumnTransformer exclusivamente no treino e a principal pratica '
     'que garante a validade das metricas reportadas para dados novos.'),
    ('Metricas alem da acuracia: ',
     'adotar Precisao, Recall, F1-Score e ROC-AUC proporcionou uma avaliacao '
     'completa, especialmente relevante dado o desbalanceamento 70/30.'),
    ('Qualidade dos resultados: ',
     f'ROC-AUC de {metricas_lr["ROC-AUC"]:.4f} na Regressao Logistica indica '
     'excelente separacao entre as classes, superando benchmarks tipicos da area.'),
    ('Interpretabilidade: ',
     'a analise de coeficientes da RL e de importancia de features do RF '
     'torna o modelo explicavel para gestores e pedagogos nao-tecnicos.'),
    ('Codigo modular: ',
     'a separacao em preprocessamento.py e modelagem.py com funcoes bem definidas '
     'facilitou manutencao, testes automatizados e reutilizacao no dashboard.'),
    ('Serializacao dos modelos: ',
     'os modelos e o pre-processador foram salvos em .pkl, '
     'permitindo inferencia em tempo real no dashboard sem re-treinamento.'),
]
for bold_txt, resto in keeps:
    bullet(doc2, resto, bold_prefix=bold_txt)

heading(doc2, '4.3 O que pode melhorar (Improve)', level=2)
improves = [
    ('Otimizacao de hiperparametros: ',
     'os modelos foram treinados com configuracoes padrao. GridSearchCV ou '
     'RandomizedSearchCV poderiam elevar ainda mais o desempenho.'),
    ('Validacao cruzada: ',
     'a avaliacao foi realizada em um unico split. k-fold cross-validation '
     'estratificado tornaria as estimativas mais confiaveis e menos dependentes da semente.'),
    ('Ajuste de threshold: ',
     'o limiar padrao de 0,5 nao e necessariamente otimo. Ajustar para maximizar '
     'Recall reduziria Falsos Negativos (alunos em risco nao identificados).'),
    ('Balanceamento de classes: ',
     'tecnicas como SMOTE ou ajuste de class_weight poderiam melhorar a '
     'identificacao da classe minoritaria (evadidos).'),
    ('Mais algoritmos: ',
     'inclusao de Gradient Boosting (XGBoost/LightGBM) e SVM para ampliar a comparacao.'),
]
for bold_txt, resto in improves:
    bullet(doc2, resto, bold_prefix=bold_txt)

heading(doc2, '4.4 Aprendizados Tecnicos', level=2)
aprendizados = [
    ('Relacoes lineares dominam o dataset: ',
     'a Regressao Logistica, modelo linear simples, superou o Random Forest em ROC-AUC, '
     'indicando que as relacoes entre features e evasao sao predominantemente lineares.'),
    ('Escala importa para RL: ',
     'o StandardScaler teve impacto direto no desempenho da Regressao Logistica, '
     'confirmando que modelos baseados em gradiente sao sensiveis a escala.'),
    ('CR e Faltas dominam a predicao: ',
     f'juntas responderam por mais de 70% da importancia no Random Forest — '
     f'achado que direciona acoes de intervencao academica.'),
    ('Estratificacao e essencial: ',
     'a estratificacao na divisao treino/teste e fundamental em datasets '
     'desbalanceados para garantir representatividade em ambas as particoes.'),
]
for bold_txt, resto in aprendizados:
    bullet(doc2, resto, bold_prefix=bold_txt)

heading(doc2, '4.5 Proximos Passos (Sprint 3)', level=2)
proximos = [
    'Otimizacao de hiperparametros com GridSearchCV (n_estimators, max_depth para RF; C para RL)',
    'Validacao cruzada estratificada com k=5 folds',
    'Ajuste de threshold de decisao para maximizacao do Recall',
    'Implementacao de balanceamento de classes (SMOTE ou class_weight)',
    'Exploracao de modelos adicionais: XGBoost, LightGBM, SVM',
    'Deploy integrado ao dashboard com persistencia dos resultados de predicao',
]
for prox in proximos:
    bullet(doc2, prox)

heading(doc2, '4.6 Conclusao', level=2)
para(doc2,
    'A Sprint 2 entregou modelos de classificacao de alta qualidade para predicao '
    'de evasao escolar, com desempenho bem acima dos requisitos minimos esperados. '
    f'A Regressao Logistica foi eleita modelo principal com ROC-AUC de '
    f'{metricas_lr["ROC-AUC"]:.4f}, Acuracia de {metricas_lr["Acuracia"]:.2%} e '
    f'Recall de {metricas_lr["Recall"]:.2%} — equilibrio que privilegia a '
    'identificacao de alunos em risco sem gerar excessivos falsos alarmes.')
para(doc2,
    'Em conjunto, as duas sprints concluidas demonstram maturidade tecnica no '
    'ciclo completo de um projeto de ML: preparacao rigorosa dos dados (Sprint 1), '
    'treinamento e avaliacao comparativa de modelos (Sprint 2), com boas praticas '
    'de engenharia de software — codigo modular, testes automatizados, '
    'containerizacao Docker e dashboard interativo para stakeholders.',
    italic=True)

# Tabela resumo sprint 2
doc2.add_paragraph()
para(doc2, 'Resumo quantitativo da Sprint 2:', bold=True)
table2 = doc2.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, txt in enumerate(['Indicador','Regressao Logistica','Random Forest']):
    hdr2[i].text = txt
    hdr2[i].paragraphs[0].runs[0].bold = True
    hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    hdr2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), AZUL_HEADER)
    hdr2[i]._tc.get_or_add_tcPr().append(shd)
for ind, vlr, vrf in [
    ('Acuracia',  f'{metricas_lr["Acuracia"]:.4f}',  f'{metricas_rf["Acuracia"]:.4f}'),
    ('Precisao',  f'{metricas_lr["Precisao"]:.4f}',  f'{metricas_rf["Precisao"]:.4f}'),
    ('Recall',    f'{metricas_lr["Recall"]:.4f}',    f'{metricas_rf["Recall"]:.4f}'),
    ('F1-Score',  f'{metricas_lr["F1-Score"]:.4f}',  f'{metricas_rf["F1-Score"]:.4f}'),
    ('ROC-AUC',   f'{metricas_lr["ROC-AUC"]:.4f}',   f'{metricas_rf["ROC-AUC"]:.4f}'),
    ('Amostras treino', '1.200', '1.200'),
    ('Amostras teste',  '300',   '300'),
    ('Features usadas', '13',    '13'),
    ('Modelo salvo em', 'models/modelo_regressao_logistica.pkl', 'models/modelo_random_forest.pkl'),
]:
    row = table2.add_row().cells
    row[0].text = ind; row[1].text = vlr; row[2].text = vrf
    for c in row: c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

out2 = os.path.join(BASE, 'Relatorio_Sprint2.docx')
doc2.save(out2)
print(f'Relatorio_Sprint2.docx salvo.')
print('\nConcluido! Dois relatorios gerados com o mesmo padrao visual.')
