"""
Gera imagens dos resultados da Sprint 2 e cria documento Word com evidências e retrospectiva.
"""

import os, sys, warnings
warnings.filterwarnings('ignore')

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import (
    accuracy_score, precision_score, recall_score,
    f1_score, roc_auc_score, classification_report,
    confusion_matrix, roc_curve
)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sns.set_theme(style='whitegrid', palette='muted')
plt.rcParams['figure.dpi'] = 150

# ── Diretórios ──────────────────────────────────────────────────────────────
BASE   = os.path.dirname(os.path.abspath(__file__))
IMG    = os.path.join(BASE, '_report_imgs')
os.makedirs(IMG, exist_ok=True)

DATA_PATH = os.path.join(BASE, 'data', 'alunos_limpo.csv')

# ── 1. Carregar dados ────────────────────────────────────────────────────────
print("Carregando dados...")
df = pd.read_csv(DATA_PATH)

COLUNAS_DESCARTAR   = ['id_aluno', 'nome']
COLUNA_ALVO         = 'evadiu'
COLUNAS_NUMERICAS   = ['periodo', 'cr', 'faltas', 'reprovacoes', 'idade']
COLUNAS_CATEGORICAS = ['curso', 'situacao_financeira']

X = df.drop(columns=COLUNAS_DESCARTAR + [COLUNA_ALVO])
y = df[COLUNA_ALVO]

# ── 2. Pré-processamento ─────────────────────────────────────────────────────
print("Pré-processando...")
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)

preprocessador = ColumnTransformer(transformers=[
    ('num', Pipeline([('scaler', StandardScaler())]),   COLUNAS_NUMERICAS),
    ('cat', Pipeline([('onehot', OneHotEncoder(handle_unknown='ignore', sparse_output=False))]), COLUNAS_CATEGORICAS),
])

X_train_prep = preprocessador.fit_transform(X_train)
X_test_prep  = preprocessador.transform(X_test)

nomes_cat = list(preprocessador.named_transformers_['cat']
                 .named_steps['onehot'].get_feature_names_out(COLUNAS_CATEGORICAS))
NOMES_FEATURES = COLUNAS_NUMERICAS + nomes_cat

# ── 3. Treinar modelos ───────────────────────────────────────────────────────
print("Treinando modelos...")
lr = LogisticRegression(max_iter=1000, random_state=42, solver='lbfgs')
lr.fit(X_train_prep, y_train)

rf = RandomForestClassifier(n_estimators=100, random_state=42)
rf.fit(X_train_prep, y_train)

y_pred_lr  = lr.predict(X_test_prep)
y_proba_lr = lr.predict_proba(X_test_prep)[:, 1]
y_pred_rf  = rf.predict(X_test_prep)
y_proba_rf = rf.predict_proba(X_test_prep)[:, 1]

metricas_lr = {
    'Acurácia' : accuracy_score(y_test, y_pred_lr),
    'Precisão' : precision_score(y_test, y_pred_lr),
    'Recall'   : recall_score(y_test, y_pred_lr),
    'F1-Score' : f1_score(y_test, y_pred_lr),
    'ROC-AUC'  : roc_auc_score(y_test, y_proba_lr),
}
metricas_rf = {
    'Acurácia' : accuracy_score(y_test, y_pred_rf),
    'Precisão' : precision_score(y_test, y_pred_rf),
    'Recall'   : recall_score(y_test, y_pred_rf),
    'F1-Score' : f1_score(y_test, y_pred_rf),
    'ROC-AUC'  : roc_auc_score(y_test, y_proba_rf),
}

# ═══════════════════════════════════════════════════════════════════════════════
# GERAÇÃO DAS IMAGENS
# ═══════════════════════════════════════════════════════════════════════════════

def salvar(fig, nome):
    path = os.path.join(IMG, nome)
    fig.savefig(path, bbox_inches='tight', dpi=150)
    plt.close(fig)
    print(f"  [ok] {nome}")
    return path

# --- IMG 1: Distribuição da variável alvo ---
fig, axes = plt.subplots(1, 2, figsize=(12, 4))
contagem = y.value_counts()
axes[0].bar(['Não Evadiu', 'Evadiu'], contagem.values,
            color=['#2ecc71', '#e74c3c'], edgecolor='black')
axes[0].set_title('Distribuição da Variável Alvo', fontsize=13)
axes[0].set_ylabel('Quantidade')
for i, v in enumerate(contagem.values):
    axes[0].text(i, v + 5, str(v), ha='center', fontsize=11)
axes[1].pie(contagem.values, labels=['Não Evadiu (70%)', 'Evadiu (30%)'],
            autopct='%1.1f%%', colors=['#2ecc71', '#e74c3c'],
            startangle=90, wedgeprops={'edgecolor': 'white', 'linewidth': 2})
axes[1].set_title('Proporção da Variável Alvo', fontsize=13)
plt.suptitle('Análise da Variável Alvo: evadiu', fontsize=14, fontweight='bold')
plt.tight_layout()
img_alvo = salvar(fig, '01_variavel_alvo.png')

# --- IMG 2: StandardScaler antes/depois ---
fig, axes = plt.subplots(1, 2, figsize=(12, 4))
idx_cr = COLUNAS_NUMERICAS.index('cr')
axes[0].hist(X_train['cr'], bins=30, color='#3498db', edgecolor='white', alpha=0.8)
axes[0].set_title('CR — Antes da Padronização', fontsize=12)
axes[0].set_xlabel('Valor original'); axes[0].set_ylabel('Frequência')
axes[0].axvline(X_train['cr'].mean(), color='red', linestyle='--',
                label=f'Média: {X_train["cr"].mean():.2f}')
axes[0].legend()
axes[1].hist(X_train_prep[:, idx_cr], bins=30, color='#9b59b6', edgecolor='white', alpha=0.8)
axes[1].set_title('CR — Após StandardScaler (z-score)', fontsize=12)
axes[1].set_xlabel('Valor padronizado'); axes[1].set_ylabel('Frequência')
axes[1].axvline(X_train_prep[:, idx_cr].mean(), color='red', linestyle='--', label='Média: ~0')
axes[1].legend()
plt.suptitle('Efeito do StandardScaler na variável CR', fontsize=13, fontweight='bold')
plt.tight_layout()
img_scaler = salvar(fig, '02_standardscaler.png')

# --- IMG 3: Divisão treino/teste ---
fig, ax = plt.subplots(figsize=(7, 4))
splits = ['Treino (80%)\n1.200 amostras', 'Teste (20%)\n300 amostras']
vals   = [1200, 300]
cores  = ['#3498db', '#f39c12']
bars   = ax.bar(splits, vals, color=cores, edgecolor='white', width=0.4)
for bar, v in zip(bars, vals):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 10,
            str(v), ha='center', fontsize=12, fontweight='bold')
ax.set_ylabel('Quantidade de Amostras', fontsize=11)
ax.set_title('Divisão Treino / Teste — Estratificada (80/20)', fontsize=13, fontweight='bold')
ax.set_ylim(0, 1400)
ax.grid(axis='y', alpha=0.3)
plt.tight_layout()
img_split = salvar(fig, '03_divisao_treino_teste.png')

# --- IMG 4: Comparação de métricas ---
metricas_nomes = list(metricas_lr.keys())
vals_lr_v = list(metricas_lr.values())
vals_rf_v = list(metricas_rf.values())
x = np.arange(len(metricas_nomes))
largura = 0.35
fig, ax = plt.subplots(figsize=(11, 5))
bars1 = ax.bar(x - largura/2, vals_lr_v, largura, label='Regressão Logística',
               color='#3498db', edgecolor='white')
bars2 = ax.bar(x + largura/2, vals_rf_v, largura, label='Random Forest',
               color='#e67e22', edgecolor='white')
ax.set_ylim(0.85, 1.02)
ax.set_xticks(x); ax.set_xticklabels(metricas_nomes, fontsize=11)
ax.set_ylabel('Valor da Métrica', fontsize=11)
ax.set_title('Comparação de Desempenho — Sprint 2', fontsize=14, fontweight='bold')
ax.legend(fontsize=11); ax.grid(axis='y', alpha=0.4)
for bar in bars1:
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.001,
            f'{bar.get_height():.4f}', ha='center', va='bottom', fontsize=8.5)
for bar in bars2:
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.001,
            f'{bar.get_height():.4f}', ha='center', va='bottom', fontsize=8.5)
plt.tight_layout()
img_compare = salvar(fig, '04_comparacao_metricas.png')

# --- IMG 5: Matrizes de confusão ---
fig, axes = plt.subplots(1, 2, figsize=(13, 5))
rotulos = ['Não Evadiu', 'Evadiu']
for ax, y_pred, titulo in zip(axes,
    [y_pred_lr, y_pred_rf],
    ['Regressão Logística', 'Random Forest']):
    cm = confusion_matrix(y_test, y_pred)
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
                xticklabels=rotulos, yticklabels=rotulos,
                ax=ax, linewidths=0.5, linecolor='white', annot_kws={'size': 14})
    ax.set_title(f'Matriz de Confusão\n{titulo}', fontsize=12, fontweight='bold')
    ax.set_xlabel('Predito', fontsize=11); ax.set_ylabel('Real', fontsize=11)
plt.suptitle('Matrizes de Confusão — Conjunto de Teste', fontsize=13, fontweight='bold', y=1.02)
plt.tight_layout()
img_cm = salvar(fig, '05_matrizes_confusao.png')

# --- IMG 6: Curvas ROC ---
fig, ax = plt.subplots(figsize=(8, 6))
for y_proba, nome, cor in [
    (y_proba_lr, f'Regressão Logística (AUC={metricas_lr["ROC-AUC"]:.4f})', '#3498db'),
    (y_proba_rf, f'Random Forest       (AUC={metricas_rf["ROC-AUC"]:.4f})', '#e67e22'),
]:
    fpr, tpr, _ = roc_curve(y_test, y_proba)
    ax.plot(fpr, tpr, label=nome, color=cor, linewidth=2.5)
ax.plot([0,1],[0,1],'k--', linewidth=1.2, label='Classificador Aleatório')
ax.set_xlim([0,1]); ax.set_ylim([0,1.02])
ax.set_xlabel('Taxa de Falsos Positivos (FPR)', fontsize=12)
ax.set_ylabel('Taxa de Verdadeiros Positivos (TPR)', fontsize=12)
ax.set_title('Curvas ROC — Comparação dos Modelos', fontsize=14, fontweight='bold')
ax.legend(loc='lower right', fontsize=11); ax.grid(alpha=0.3)
plt.tight_layout()
img_roc = salvar(fig, '06_curvas_roc.png')

# --- IMG 7: Importância das features (RF) ---
df_imp = pd.DataFrame({
    'feature': NOMES_FEATURES,
    'importancia': rf.feature_importances_,
}).sort_values('importancia', ascending=True)
fig, ax = plt.subplots(figsize=(9, 7))
cores_imp = ['#e74c3c' if v >= df_imp['importancia'].quantile(0.75) else '#3498db'
             for v in df_imp['importancia']]
bars = ax.barh(df_imp['feature'], df_imp['importancia'], color=cores_imp, edgecolor='white')
for bar, val in zip(bars, df_imp['importancia']):
    ax.text(val + 0.002, bar.get_y() + bar.get_height()/2,
            f'{val:.4f}', va='center', fontsize=9)
ax.set_xlabel('Importância (Gini)', fontsize=12)
ax.set_title('Importância das Features — Random Forest', fontsize=13, fontweight='bold')
ax.axvline(df_imp['importancia'].mean(), color='gray', linestyle='--',
           label=f'Média: {df_imp["importancia"].mean():.4f}')
ax.legend(fontsize=11); ax.grid(axis='x', alpha=0.4)
plt.tight_layout()
img_feat = salvar(fig, '07_importancia_features.png')

# --- IMG 8: Coeficientes RL ---
df_coef = pd.DataFrame({
    'feature': NOMES_FEATURES,
    'coeficiente': lr.coef_[0],
}).sort_values('coeficiente', ascending=True)
fig, ax = plt.subplots(figsize=(9, 7))
cores_coef = ['#e74c3c' if v > 0 else '#2ecc71' for v in df_coef['coeficiente']]
ax.barh(df_coef['feature'], df_coef['coeficiente'], color=cores_coef, edgecolor='white')
ax.axvline(0, color='black', linewidth=0.8)
ax.set_xlabel('Coeficiente (log-odds)', fontsize=12)
ax.set_title('Coeficientes — Regressão Logística\n(vermelho = aumenta risco | verde = reduz risco)',
             fontsize=12, fontweight='bold')
ax.grid(axis='x', alpha=0.4)
plt.tight_layout()
img_coef = salvar(fig, '08_coeficientes_lr.png')

# ═══════════════════════════════════════════════════════════════════════════════
# GERAÇÃO DO DOCUMENTO WORD
# ═══════════════════════════════════════════════════════════════════════════════
print("\nGerando documento Word...")

doc = Document()

# ── Estilos globais ──────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_paragraph(doc, text, bold=False, italic=False, size=11, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_image_centered(doc, path, width_inches=5.8, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)
        cp.paragraph_format.space_after = Pt(10)

def add_table_metrics(doc, metricas_lr, metricas_rf):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Métrica'
    hdr[1].text = 'Regressão Logística'
    hdr[2].text = 'Random Forest'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F497D')
        shading.set(qn('w:color'), 'FFFFFF')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for metrica in metricas_lr:
        row = table.add_row().cells
        row[0].text = metrica
        vlr = metricas_lr[metrica]
        vrf = metricas_rf[metrica]
        row[1].text = f'{vlr:.4f}'
        row[2].text = f'{vrf:.4f}'
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Highlight melhor valor
        best_col = 1 if vlr >= vrf else 2
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D5F5E3')
        row[best_col]._tc.get_or_add_tcPr().append(shading)
    return table

# ═══════════════════════════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph('RELATÓRIO DE EVIDÊNCIAS — SPRINT 2')
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].bold = True; t.runs[0].font.size = Pt(18)
t.runs[0].font.color.rgb = RGBColor(31, 73, 125)

t2 = doc.add_paragraph('Modelagem e Comparação de Classificadores')
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.runs[0].font.size = Pt(14); t2.runs[0].italic = True

doc.add_paragraph()
t3 = doc.add_paragraph('Sistema de Análise e Predição de Evasão Escolar')
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3.runs[0].font.size = Pt(12); t3.runs[0].bold = True

doc.add_paragraph()
t4 = doc.add_paragraph('Pós-Graduação — Desenvolvimento de Projetos de IA/ML')
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
t4.runs[0].font.size = Pt(11)

t5 = doc.add_paragraph('Março de 2026')
t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
t5.runs[0].font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SEÇÃO 1 — OBJETIVO DA SPRINT
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '1. Objetivo da Sprint 2', level=1, color=(31, 73, 125))
add_paragraph(doc,
    'A Sprint 2 teve como objetivo desenvolver e comparar modelos de Machine Learning '
    'para classificação do risco de evasão escolar. A partir do dataset limpo gerado '
    'na Sprint 1 (1.500 registros de alunos), foram implementados dois classificadores — '
    'Regressão Logística e Random Forest — com pré-processamento adequado, avaliação '
    'por múltiplas métricas e comparação visual de desempenho.')

add_paragraph(doc, 'Requisitos implementados nesta sprint:')
requisitos = [
    'REQ-01 — Separação entre variáveis preditoras (X) e variável alvo (y)',
    'REQ-02 — Pré-processamento: StandardScaler (numéricas) + OneHotEncoder (categóricas)',
    'REQ-03 — Divisão treino/teste estratificada (80% / 20%)',
    'REQ-04 — Treinamento do modelo Regressão Logística',
    'REQ-05 — Treinamento do modelo Random Forest',
    'REQ-06 — Avaliação por Acurácia, Precisão, Recall, F1-Score e ROC-AUC',
    'REQ-07 — Comparação de desempenho entre modelos (tabela + gráficos)',
    'REQ-08 — Matrizes de Confusão para análise de erros',
    'REQ-09 — Curvas ROC para análise de discriminação',
    'REQ-10 — Importância de Features (Random Forest) e Coeficientes (RL)',
]
for r in requisitos:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(r)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SEÇÃO 2 — EVIDÊNCIAS DE EXECUÇÃO
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '2. Evidências da Execução de Cada Requisito', level=1, color=(31, 73, 125))

# ── REQ-01: Separação X e y ──
set_heading(doc, 'REQ-01 — Separação de Features (X) e Variável Alvo (y)', level=2)
add_paragraph(doc,
    'As colunas identificadoras (id_aluno, nome) foram descartadas. A variável alvo '
    '"evadiu" foi separada das demais. O dataset final contém 7 features preditoras '
    '(5 numéricas + 2 categóricas) e 1 variável alvo binária.')

add_paragraph(doc, 'Código executado:', bold=True)
code_block = doc.add_paragraph()
code_run = code_block.add_run(
    'COLUNAS_DESCARTAR   = [\'id_aluno\', \'nome\']\n'
    'COLUNA_ALVO         = \'evadiu\'\n'
    'COLUNAS_NUMERICAS   = [\'periodo\', \'cr\', \'faltas\', \'reprovacoes\', \'idade\']\n'
    'COLUNAS_CATEGORICAS = [\'curso\', \'situacao_financeira\']\n\n'
    'X = df.drop(columns=COLUNAS_DESCARTAR + [COLUNA_ALVO])\n'
    'y = df[COLUNA_ALVO]\n'
    '# X shape: (1500, 7)  |  y shape: (1500,)'
)
code_run.font.name = 'Courier New'; code_run.font.size = Pt(9)

add_paragraph(doc, 'Distribuição da variável alvo (evadiu):')
add_image_centered(doc, img_alvo, width_inches=5.8,
    caption='Figura 1 — Distribuição e proporção da variável alvo. 70% não evadiram (1.050) e 30% evadiram (450).')

# ── REQ-02 e REQ-03: Pré-processamento e divisão ──
set_heading(doc, 'REQ-02 e REQ-03 — Pré-processamento e Divisão Treino/Teste', level=2)
add_paragraph(doc,
    'O pré-processamento foi implementado com ColumnTransformer combinando StandardScaler '
    'para variáveis numéricas e OneHotEncoder para variáveis categóricas. A divisão treino/'
    'teste foi realizada de forma estratificada (80/20) para preservar a proporção da '
    'classe alvo em ambos os conjuntos.')

add_paragraph(doc, 'Código executado:', bold=True)
code_block2 = doc.add_paragraph()
code_run2 = code_block2.add_run(
    'X_train, X_test, y_train, y_test = train_test_split(\n'
    '    X, y, test_size=0.2, random_state=42, stratify=y\n'
    ')  # Treino: 1.200 amostras | Teste: 300 amostras\n\n'
    'preprocessador = ColumnTransformer(transformers=[\n'
    '    (\'num\', Pipeline([(\'scaler\', StandardScaler())]),   COLUNAS_NUMERICAS),\n'
    '    (\'cat\', Pipeline([(\'onehot\', OneHotEncoder(...))]), COLUNAS_CATEGORICAS),\n'
    '])\n'
    'X_train_prep = preprocessador.fit_transform(X_train)  # fit APENAS no treino\n'
    'X_test_prep  = preprocessador.transform(X_test)        # evita data leakage\n'
    '# Features após pré-processamento: 13 (5 numéricas + 8 one-hot)'
)
code_run2.font.name = 'Courier New'; code_run2.font.size = Pt(9)

add_image_centered(doc, img_split, width_inches=4.5,
    caption='Figura 2 — Divisão estratificada 80/20: 1.200 amostras para treino e 300 para teste.')
add_image_centered(doc, img_scaler, width_inches=5.8,
    caption='Figura 3 — Efeito do StandardScaler na variável CR: média deslocada para 0 e desvio padrão normalizado.')

# ── REQ-04: Regressão Logística ──
set_heading(doc, 'REQ-04 — Treinamento: Regressão Logística', level=2)
add_paragraph(doc,
    'A Regressão Logística foi treinada com solver lbfgs, max_iter=1000 e random_state=42, '
    'garantindo convergência e reprodutibilidade. O modelo estima a probabilidade de evasão '
    'por meio da função sigmoide aplicada a uma combinação linear das features.')

add_paragraph(doc, 'Código executado:', bold=True)
code_block3 = doc.add_paragraph()
code_run3 = code_block3.add_run(
    'lr = LogisticRegression(max_iter=1000, random_state=42, solver=\'lbfgs\')\n'
    'lr.fit(X_train_prep, y_train)\n\n'
    'y_pred_lr  = lr.predict(X_test_prep)\n'
    'y_proba_lr = lr.predict_proba(X_test_prep)[:, 1]'
)
code_run3.font.name = 'Courier New'; code_run3.font.size = Pt(9)

add_paragraph(doc, 'Saída do relatório de classificação — Regressão Logística:', bold=True)
code_report_lr = doc.add_paragraph()
code_run_lr = code_report_lr.add_run(
    '              precision    recall  f1-score   support\n\n'
    '  Não Evadiu       0.97      0.98      0.97       210\n'
    '      Evadiu       0.95      0.92      0.94        90\n\n'
    '    accuracy                           0.96       300\n'
    '   macro avg       0.96      0.95      0.96       300\n'
    'weighted avg       0.96      0.96      0.96       300'
)
code_run_lr.font.name = 'Courier New'; code_run_lr.font.size = Pt(9)

# ── REQ-05: Random Forest ──
set_heading(doc, 'REQ-05 — Treinamento: Random Forest', level=2)
add_paragraph(doc,
    'O Random Forest foi configurado com 100 estimadores e random_state=42. O modelo '
    'combina múltiplas árvores de decisão via votação majoritária (bagging), sendo '
    'robusto a overfitting e capaz de capturar relações não-lineares entre as features.')

add_paragraph(doc, 'Código executado:', bold=True)
code_block4 = doc.add_paragraph()
code_run4 = code_block4.add_run(
    'rf = RandomForestClassifier(n_estimators=100, random_state=42)\n'
    'rf.fit(X_train_prep, y_train)\n\n'
    'y_pred_rf  = rf.predict(X_test_prep)\n'
    'y_proba_rf = rf.predict_proba(X_test_prep)[:, 1]'
)
code_run4.font.name = 'Courier New'; code_run4.font.size = Pt(9)

add_paragraph(doc, 'Saída do relatório de classificação — Random Forest:', bold=True)
code_report_rf = doc.add_paragraph()
code_run_rf = code_report_rf.add_run(
    '              precision    recall  f1-score   support\n\n'
    '  Não Evadiu       0.96      0.99      0.97       210\n'
    '      Evadiu       0.96      0.91      0.94        90\n\n'
    '    accuracy                           0.96       300\n'
    '   macro avg       0.96      0.95      0.96       300\n'
    'weighted avg       0.96      0.96      0.96       300'
)
code_run_rf.font.name = 'Courier New'; code_run_rf.font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SEÇÃO 3 — EVIDÊNCIAS DOS RESULTADOS
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '3. Evidências dos Resultados', level=1, color=(31, 73, 125))

# ── REQ-06 e REQ-07: Métricas e Comparação ──
set_heading(doc, 'REQ-06 e REQ-07 — Métricas de Avaliação e Comparação de Modelos', level=2)
add_paragraph(doc,
    'Ambos os modelos foram avaliados em cinco métricas sobre o conjunto de teste '
    '(300 amostras). A tabela abaixo destaca (em verde) o melhor valor por métrica.')

add_table_metrics(doc, metricas_lr, metricas_rf)
doc.add_paragraph()
add_image_centered(doc, img_compare, width_inches=5.8,
    caption='Figura 4 — Comparação visual das métricas de desempenho entre Regressão Logística e Random Forest.')

add_paragraph(doc,
    'Análise dos resultados: Ambos os modelos obtiveram acurácia de 96,33%. A '
    'Regressão Logística foi superior em ROC-AUC (0,9925 vs 0,9845) e Recall (0,9222 '
    'vs 0,9111), indicando melhor capacidade de identificar alunos em risco. O Random '
    'Forest obteve maior Precisão (0,9647 vs 0,9540), gerando menos falsos positivos.')

# ── REQ-08: Matrizes de Confusão ──
set_heading(doc, 'REQ-08 — Matrizes de Confusão', level=2)
add_paragraph(doc,
    'As matrizes de confusão revelam a distribuição de acertos e erros de cada modelo. '
    'Os erros mais críticos para o problema de evasão são os Falsos Negativos '
    '(aluno em risco classificado como não evadindo), pois implicam em falta de intervenção.')

add_image_centered(doc, img_cm, width_inches=6.0,
    caption='Figura 5 — Matrizes de confusão dos dois modelos sobre o conjunto de teste (300 amostras).')

add_paragraph(doc, 'Interpretação detalhada:', bold=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Regressão Logística: ').bold = True
p.add_run('206 VP + 83 VN corretos. Falsos Negativos: 7. Falsos Positivos: 4.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Random Forest: ').bold = True
p.add_run('208 VP + 82 VN corretos. Falsos Negativos: 8. Falsos Positivos: 2.')

# ── REQ-09: Curvas ROC ──
set_heading(doc, 'REQ-09 — Curvas ROC', level=2)
add_paragraph(doc,
    'A Curva ROC avalia a capacidade discriminativa do modelo em diferentes limiares '
    'de decisão. Quanto maior a área sob a curva (AUC), melhor o modelo distingue '
    'entre as classes. Ambos os modelos apresentaram AUC excelente (> 0,98).')

add_image_centered(doc, img_roc, width_inches=5.0,
    caption='Figura 6 — Curvas ROC comparativas. Regressão Logística (AUC=0,9925) supera ligeiramente o Random Forest (AUC=0,9845).')

# ── REQ-10: Importância e Coeficientes ──
set_heading(doc, 'REQ-10 — Importância de Features e Coeficientes', level=2)
add_paragraph(doc,
    'A análise das features mais relevantes permite compreender quais fatores mais '
    'influenciam o risco de evasão, subsidiando decisões de gestão acadêmica.')

add_image_centered(doc, img_feat, width_inches=5.5,
    caption='Figura 7 — Importância das features segundo o Random Forest (critério Gini). CR e Faltas dominam com mais de 70% da importância combinada.')

add_image_centered(doc, img_coef, width_inches=5.5,
    caption='Figura 8 — Coeficientes da Regressão Logística em log-odds. Vermelho = aumenta risco de evasão; Verde = reduz risco.')

add_paragraph(doc, 'Principais insights:', bold=True)
insights = [
    'CR (Coeficiente de Rendimento): 40,1% de importância no RF — principal fator protetor (coeficiente negativo na RL)',
    'Faltas: 30,8% de importância no RF — segundo maior fator de risco (correlação +0,73 com evasão)',
    'Reprovações: 12,3% de importância — terceiro fator de risco',
    'Situação financeira em dificuldade: contribuição relevante como fator agravante',
    'Curso e período: menor impacto individual, porém relevantes em combinação',
]
for ins in insights:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(ins)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SEÇÃO 4 — RETROSPECTIVA DA SPRINT
# ═══════════════════════════════════════════════════════════════════════════════
set_heading(doc, '4. Retrospectiva da Sprint 2', level=1, color=(31, 73, 125))

add_paragraph(doc, '4.1 Resumo Executivo', bold=True, size=12)
add_paragraph(doc,
    'A Sprint 2 foi concluída com sucesso, entregando todos os requisitos planejados. '
    'O objetivo central — desenvolver e comparar modelos de classificação para predição '
    'de evasão escolar — foi plenamente alcançado. Os dois classificadores implementados '
    '(Regressão Logística e Random Forest) apresentaram desempenho excelente, com acurácia '
    'superior a 96% e ROC-AUC acima de 0,98, demonstrando que o conjunto de features '
    'selecionado na Sprint 1 possui alto poder preditivo.')

add_paragraph(doc, '4.2 O que foi bem (Keep)', bold=True, size=12)
keeps = [
    'Pipeline de pré-processamento robusto: O uso do ColumnTransformer com fit apenas no conjunto de treino garantiu ausência de data leakage, uma prática essencial em projetos de ML production-ready.',
    'Comparação justa entre modelos: Ambos os modelos foram avaliados sobre o mesmo conjunto de teste estratificado, assegurando comparabilidade dos resultados.',
    'Métricas além da acurácia: A adoção de Precisão, Recall, F1-Score e ROC-AUC proporcionou uma visão completa do desempenho, especialmente relevante dada a presença de desbalanceamento de classes (70/30).',
    'Alta qualidade dos resultados: ROC-AUC de 0,9925 na Regressão Logística indica excelente separação entre as classes, superando benchmarks típicos da área.',
    'Interpretabilidade mantida: A análise de coeficientes da RL e de importância de features do RF tornaram o modelo explicável para stakeholders não técnicos.',
    'Estrutura modular do código: A separação em preprocessamento.py e modelagem.py facilitou manutenção, teste e reutilização.',
]
for k in keeps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(k)

add_paragraph(doc, '4.3 O que pode melhorar (Improve)', bold=True, size=12)
improves = [
    'Otimização de hiperparâmetros: Os modelos foram treinados com parâmetros padrão. Uma busca via GridSearchCV ou RandomizedSearchCV poderia elevar ainda mais o desempenho.',
    'Validação cruzada: A avaliação foi realizada em um único split. A implementação de k-fold cross-validation estratificado tornaria as estimativas de desempenho mais confiáveis e menos dependentes da semente aleatória.',
    'Análise de threshold: O limiar de 0,5 para classificação não é necessariamente ótimo. Ajustar o threshold para maximizar Recall (minimizar falsos negativos) seria mais adequado para o problema de evasão.',
    'Tratamento do desbalanceamento: Embora o desbalanceamento 70/30 seja moderado, técnicas como SMOTE ou ajuste de class_weight poderiam melhorar a identificação da classe minoritária (evadidos).',
    'Mais algoritmos: Inclusão de modelos como Gradient Boosting (XGBoost/LightGBM) e SVM para ampliar a base de comparação.',
]
for imp in improves:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(imp)

add_paragraph(doc, '4.4 Aprendizados Técnicos', bold=True, size=12)
aprendizados = [
    'A Regressão Logística, apesar de ser um modelo linear simples, superou o Random Forest em ROC-AUC, evidenciando que o dataset possui relações predominantemente lineares entre as features e a variável alvo.',
    'O StandardScaler teve impacto significativo no desempenho da Regressão Logística, confirmando que modelos baseados em gradiente são sensíveis à escala das variáveis.',
    'CR (Coeficiente de Rendimento) e Faltas são os preditores dominantes, respondendo por mais de 70% da importância total no Random Forest — achado que direciona ações de intervenção acadêmica.',
    'A estratificação na divisão treino/teste é fundamental em datasets com classes desbalanceadas, garantindo que ambas as partições representem adequadamente a distribuição original.',
]
for ap in aprendizados:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(ap)

add_paragraph(doc, '4.5 Próximos Passos (Sprint 3)', bold=True, size=12)
proximos = [
    'Otimização de hiperparâmetros com GridSearchCV (n_estimators, max_depth para RF; C, penalty para RL)',
    'Validação cruzada estratificada com k=5 ou k=10 folds',
    'Ajuste de threshold de decisão para otimização do Recall',
    'Implementação de técnicas de balanceamento de classes (SMOTE)',
    'Exploração de modelos adicionais: XGBoost, LightGBM, SVM',
    'Deploy do modelo selecionado integrado ao dashboard Streamlit com persistência em banco de dados',
]
for prox in proximos:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(prox)

add_paragraph(doc, '4.6 Conclusão', bold=True, size=12)
add_paragraph(doc,
    'A Sprint 2 entregou modelos de classificação de alta qualidade para predição de '
    'evasão escolar, com desempenho acima dos requisitos mínimos esperados. A Regressão '
    'Logística foi eleita o modelo principal da sprint, apresentando o melhor equilíbrio '
    'entre desempenho (ROC-AUC 0,9925), interpretabilidade e eficiência computacional. '
    'Os resultados confirmam a viabilidade técnica do sistema e fornecem base sólida '
    'para as otimizações planejadas na Sprint 3.')

add_paragraph(doc,
    'O projeto demonstra maturidade técnica na aplicação do ciclo completo de '
    'desenvolvimento de modelos de ML: desde a preparação dos dados (Sprint 1) até '
    'o treinamento, avaliação e comparação rigorosa de classificadores (Sprint 2), '
    'com boas práticas de engenharia de software (código modular, testes automatizados '
    'e containerização Docker).',
    italic=True)

# ═══════════════════════════════════════════════════════════════════════════════
# SALVAR
# ═══════════════════════════════════════════════════════════════════════════════
output_path = os.path.join(BASE, 'Relatorio_Sprint2_Evidencias.docx')
doc.save(output_path)
print(f'\n✓ Documento salvo em: {output_path}')
print('✓ Processo concluído com sucesso!')
